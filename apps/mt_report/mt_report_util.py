#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.
import json
import logging.config
import os
import time

from docx import Document
from docx.shared import RGBColor, Cm
from langchain_core.prompts import ChatPromptTemplate

from common import statistic_util
from common.agt_util import get_model
from common.cm_utils import extract_md_content, rmv_think_block, estimate_tokens
from common.sys_init import init_yml_cfg

log_config_path = 'logging.conf'
if os.path.exists(log_config_path):
    logging.config.fileConfig(log_config_path, encoding="utf-8")
else:
    from common.const import LOG_FORMATTER
    logging.basicConfig(level=logging.INFO,format= LOG_FORMATTER, force=True)
logger = logging.getLogger(__name__)

EXTRACT_ABS_PROMPT = """
# 角色与任务
你是一个AI助手，专门负责将杂乱的会议记录草稿，按照预设的字段提取和整理信息。

# 输入信息
## 【原始会议记录草稿】
{input_txt}

## 【待填充字段定义】
这是一个Word模板中需要填充的字段列表。每个字段编号对应一个需要从上述草稿中提取信息的主题。
{field_dict}

# 任务指令
你的任务是根据【原始会议记录草稿】，为【待填充字段定义】中的每一个字段，找到或总结出最匹配的文本内容。

# 输出要求
1.  **严格按格式输出**：你必须输出一个JSON对象，其键是字段编号（如 "1", "2"），其值是你为该字段从草稿中提取或总结的文本。
2.  **内容处理原则**：
 - **提取优先**：如果草稿中有明确对应字段的原文，请直接提取。
 - **总结为辅**：如果某个字段的信息在草稿中分散在多处，请用简洁、客观的语言进行归纳总结。
 - **忠于原意**：切勿添加草稿中不存在的信息或做出主观臆断。如果某个字段在草稿中找不到对应信息，该字段的值为空字符串 ""。
3.  **特别关注**：对于“下一步行动项”这类字段，请确保提取出“任务内容”、“负责人”、“截止时间”等关键要素（如果草稿中有的话）。

# 输出示例
你的输出应该看起来像这样：
{{
"1": "关于Q3产品发布计划的评审会议",
"2": "2023年10月26日下午2点",
"3": "张三、李四、王五",
"4": "评审了产品原型，决议按方案B进行开发，并增加用户反馈模块。",
"5": "1. 完成最终版UI设计图 - 李四 - 11月3日； 2. 准备服务器部署方案 - 王五 - 11月10日"
}}

现在，请开始处理上述提供的【原始会议记录草稿】和【待填充字段定义】，并输出JSON结果。
 """

GET_FIELD_PROMPT = """
你是一个文档处理专家，需要分析一个Word文档模板的段落结构。

【任务】
请仔细阅读以下从Word模板中提取出的段落列表。你的目标是识别出哪些段落是“文档的字段标题”。

【“字段标题”的定义】
- 它是一个**引导句或小节标题**，其后的内容需要根据每次会议的具体信息来填充。
- 它通常是**一个名词性短语或一个不完整的句子**，后面跟着冒号、空格或者直接留空。
- 例如：“会议主题：", “出席人员：", “主要决议", “下一步计划” 都是典型的字段标题。
- 而具体的填写内容、大段的说明文字、页脚页码等，则**不是**字段标题。

【输入数据】
以下是模板的段落列表，格式为 [索引号]: 段落内容
{all_paragraphs}

【输出要求】
请严格按照以下JSON格式输出结果，且只输出JSON：
{{
  "field_dict": {{
    "段落索引号": "你识别出的字段标题名（请精简概括）",
    ...
  }}
}}

请只输出你确信是“字段标题”的段落。如果一个段落看起来不重要或不是字段标题，请忽略它。

"""

FILED_DICT = {
    1: "会议主题",
    2: "会议时间",
    3: "出席人员",
    4: "主要议程与决议",
    5: "下一步行动项"
}


def get_txt_abs(uid: int, cfg: dict, input_txt: str, field_dict: dict) -> dict:
    """
    从文本中抽取关键信息，输出关键信息词典
    """
    max_retries = 3
    backoff_times = [5, 10, 20]
    model = None
    for attempt in range(max_retries + 1):
        try:
            if attempt > 0:
                wait_time = backoff_times[attempt - 1]
                logger.info(f"retry #{attempt} times after {wait_time}s")
                time.sleep(wait_time)
            prompt = ChatPromptTemplate.from_template(EXTRACT_ABS_PROMPT)
            model = get_model(cfg)
            chain = prompt | model
            arg_dict = {"input_txt": input_txt, "field_dict": field_dict}
            input_text = prompt.format(input_txt=input_txt, field_dict=field_dict)
            # logger.info(f"{uid}, start_calc_txt_token, {input_text}")
            input_tokens = estimate_tokens(input_text)
            logger.info(f"{uid}, input_tokens, {input_tokens}")
            statistic_util.add_input_token_by_uid(uid, input_tokens)
            logger.info(f"submit_arg_dict_to_llm, [{arg_dict}], llm[{cfg['api']}]")
            response = chain.invoke(arg_dict)
            output_txt = extract_md_content(rmv_think_block(response.content), "json")
            output_tokens = estimate_tokens(response.content)
            logger.info(f"{uid}, output_tokens, {output_tokens}")
            statistic_util.add_output_token_by_uid(uid, output_tokens)
            dispose(model)
            return json.loads(output_txt)
        except Exception as ex:
            last_exception = ex
            logger.error(f"retry_failed, retry_time={attempt}, {str(ex)}")
            if attempt < max_retries:
                continue
            dispose(model)
            logger.error(f"all_retries_exhausted, {input_txt}, {field_dict}")
            raise last_exception
    return {}

def get_template_field(uid: int, cfg: dict, all_paragraphs: str) -> dict:
    """
    从 Word 模板的文本中获取需要填充的核心要素词典
    """
    max_retries = 3
    backoff_times = [5, 10, 20]
    model = None
    for attempt in range(max_retries + 1):
        try:
            if attempt > 0:
                wait_time = backoff_times[attempt - 1]
                logger.info(f"retry #{attempt} times after {wait_time}s")
                time.sleep(wait_time)
            prompt = ChatPromptTemplate.from_template(GET_FIELD_PROMPT)
            model = get_model(cfg)
            chain = prompt | model
            arg_dict = {"all_paragraphs": all_paragraphs}
            input_text = prompt.format(all_paragraphs=all_paragraphs)
            input_tokens = estimate_tokens(input_text)
            logger.info(f"{uid}, input_tokens, {input_tokens}")
            statistic_util.add_input_token_by_uid(uid, input_tokens)
            logger.info(f"submit_arg_dict_to_llm, [{arg_dict}], llm[{cfg['api']}]")
            response = chain.invoke(arg_dict)
            output_txt = extract_md_content(rmv_think_block(response.content), "json")
            out_tokens = estimate_tokens(response.content)
            logger.info(f"{uid}, out_tokens, {out_tokens}")
            statistic_util.add_output_token_by_uid(uid, out_tokens)
            dispose(model)
            return json.loads(output_txt)
        except Exception as ex:
            last_exception = ex
            logger.error(f"retry_failed, retry_time={attempt}, {str(ex)}")
            if attempt < max_retries:
                continue
            dispose(model)
            logger.error(f"all_retries_exhausted, {all_paragraphs}")
            raise last_exception
    return {}

def get_doc_content(input_file: str):
    """
    返回 docx 文档的全文内容，格式为
    [索引号]: 段落内容
    """
    doc = Document(input_file)
    content = []
    for index, para in enumerate(doc.paragraphs):
        para_txt = f"[{index}]: {para.text}"
        content.append(para_txt)
    return "\n".join(content)



def insert_para_to_doc(input_doc: str, output_doc: str, para_dict: dict[str, str]) -> bool:
    """
    :param input_doc: 输入文档的路径
    :param output_doc: 输出文档的路径
    :param para_dict: 段落ID和段落之后需要插入的文本对应关系的字典
    :return: 操作是否成功
    """
    try:
        doc = Document(input_doc)
        paragraphs = doc.paragraphs

        if not para_dict:
            logger.warning(f"没有有效的插入文本, 保存原始文档:{input_doc} 至 {output_doc}")
            doc.save(output_doc)
            return True

        # 将字典键转换为整数并排序（从后往前处理避免索引变化）
        sorted_indices = sorted([int(idx) for idx in para_dict.keys()], reverse=True)

        # 从后往前插入，避免索引变化
        for para_index in sorted_indices:
            if para_index >= len(paragraphs):
                logger.warning(
                    f"段落索引 {para_index} 超出范围（总段落数：{len(paragraphs)}），跳过, input_file={input_doc}")
                continue
            original_para = paragraphs[para_index]
            generated_text = para_dict[str(para_index)]  # 获取对应的文本内容

            # 创建新段落并插入
            new_para = doc.add_paragraph()
            new_para.paragraph_format.first_line_indent = Cm(1)  # 调整为更合理的缩进值
            new_para.paragraph_format.space_after = Cm(0.35)  # 添加段后间距
            new_run = new_para.add_run(generated_text)
            new_run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色文本

            # 插入到原始段落后面
            original_para._p.addnext(new_para._p)
            logger.debug(f"已在段落 {para_index} 后插入生成内容: {generated_text[:50]}...")

        doc.save(output_doc)
        logger.info(f"使用段落索引更新文档成功: {output_doc}")
        return True
    except Exception as e:
        logger.exception(f"使用段落索引更新文档失败: input_file={input_doc}, output_doc={output_doc}")
        return False


def dispose(model):
    if 'model' in locals():
        del model
        import torch
        torch.cuda.empty_cache()


if __name__ =="__main__":
    my_cfg = init_yml_cfg()
    input_doc = "/home/rd/Desktop/test_template.docx"
    report_txt_file = "/home/rd/Desktop/report.txt"
    doc_txt = get_doc_content(input_doc)
    logger.info(f"doc_txt={doc_txt}")
    doc_field = get_template_field(my_cfg, doc_txt)
    logger.info(f"doc_field={doc_field}")
    with open(report_txt_file, 'r', encoding='utf-8') as file:
        input_txt = file.read()
    logger.info(f"input_txt={input_txt}")
    txt_abs = get_txt_abs(my_cfg, input_txt, doc_field.get("field_dict"))
    logger.info(f"txt_abs={txt_abs}")
    output_doc = "/home/rd/Desktop/output.docx"
    insert_para_to_doc(input_doc, output_doc, txt_abs)