#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.

import json
import sys
import os
import jwt
import logging.config
import sqlite3
from datetime import datetime

from dotenv import load_dotenv
from flask import Flask, render_template, send_from_directory, abort, request, jsonify, g

from apps.online_office.office_util import generate_jwt_token, JWT_SECRET, get_content_type, generate_onlyoffice_config, \
    get_docker_host, get_file_type
from common.const import UPLOAD_FOLDER
from common.i18n._hooks import register_i18n
from common.i18n import get_msg

import uuid
import shutil

# 在已有配置后添加文档相关的配置
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCUMENT_FOLDER = os.path.join(BASE_DIR, 'documents')
DOCUMENT_UPLOAD_FOLDER = os.path.join(DOCUMENT_FOLDER, 'uploads')
DOCUMENT_PREVIEW_FOLDER = os.path.join(DOCUMENT_FOLDER, 'preview')
DOCUMENT_TEMP_FOLDER = os.path.join(DOCUMENT_FOLDER, 'temp')

# 创建必要的文件夹
for folder in [DOCUMENT_FOLDER, DOCUMENT_UPLOAD_FOLDER, DOCUMENT_PREVIEW_FOLDER, DOCUMENT_TEMP_FOLDER]:
    os.makedirs(folder, exist_ok=True)

# 扩展允许的文件类型
ALLOWED_DOC_EXTENSIONS = {'docx', 'doc', 'txt', 'pdf', 'xlsx', 'pptx'}
DOC_MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# 数据库配置
DATABASE = os.path.join(os.path.dirname(__file__), 'documents.db')

# 初始化数据库
def init_database():
    """初始化SQLite数据库"""
    with sqlite3.connect(DATABASE) as conn:
        cursor = conn.cursor()
        
        # 创建文档表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id TEXT PRIMARY KEY,
                original_filename TEXT NOT NULL,
                filename TEXT NOT NULL,
                file_type TEXT NOT NULL,
                file_ext TEXT NOT NULL,
                size INTEGER NOT NULL,
                upload_time TEXT NOT NULL,
                url TEXT NOT NULL,
                path TEXT NOT NULL,
                key TEXT NOT NULL,
                user_id TEXT DEFAULT 'anonymous',
                status TEXT DEFAULT 'active',
                last_modified TEXT,
                analysis_time TEXT
            )
        ''')
        
        # 创建文档分析结果表
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS document_analysis (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                doc_id TEXT NOT NULL,
                category TEXT NOT NULL,
                severity TEXT NOT NULL,
                position TEXT NOT NULL,
                description TEXT NOT NULL,
                suggestion TEXT NOT NULL,
                status TEXT DEFAULT 'pending',  -- pending, accepted, ignored
                created_time TEXT NOT NULL,
                FOREIGN KEY (doc_id) REFERENCES documents (id) ON DELETE CASCADE
            )
        ''')
        
        conn.commit()

# 初始化数据库
init_database()

# 获取数据库连接
def get_db():
    """获取数据库连接"""
    if 'db' not in g:
        g.db = sqlite3.connect(DATABASE)
        g.db.row_factory = sqlite3.Row
    return g.db

# 关闭数据库连接
@app.teardown_appcontext
def close_db(error):
    """关闭数据库连接"""
    db = g.pop('db', None)
    if db is not None:
        db.close()


ONLY_OFFICE_API = "http://localhost"

# 创建 Flask 应用
app = Flask(__name__, static_folder=None)
register_i18n(app, scope="online_office")

# 错误处理
@app.errorhandler(404)
def not_found_error(error):
    logger.error(f"404错误: {error}")
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'error': get_msg('online_office.resource_not_found')}), 404
    return render_template('error.html', error_code=404, error_message=get_msg('online_office.page_not_found')), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"500错误: {error}")
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'error': get_msg('online_office.server_internal_error')}), 500
    return render_template('error.html', error_code=500, error_message=get_msg('online_office.server_internal_error')), 500

@app.errorhandler(413)
def request_entity_too_large(error):
    logger.error(f"413错误: 文件太大")
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'error': get_msg('online_office.file_too_large_max', max=50)}), 413
    return render_template('error.html', error_code=413, error_message=get_msg('online_office.file_too_large_short')), 413

@app.errorhandler(400)
def bad_request_error(error):
    logger.error(f"400错误: {error}")
    if request.path.startswith('/api/'):
        return jsonify({'success': False, 'error': get_msg('online_office.bad_request_format')}), 400
    return render_template('error.html', error_code=400, error_message=get_msg('online_office.bad_request_format')), 400

log_config_path = 'logging.conf'
if os.path.exists(log_config_path):
    logging.config.fileConfig(log_config_path, encoding="utf-8", disable_existing_loggers=False)
    print(f"使用日志配置文件: {log_config_path}")
else:
    print("日志配置文件不存在，使用默认配置")
    from common.const import LOG_FORMATTER
    logging.basicConfig(level=logging.INFO,format=LOG_FORMATTER,force=True,stream=sys.stdout)

logger = logging.getLogger(__name__)
logger.info("应用程序启动")

# 加载环境变量
load_dotenv()

# 确保上传文件夹存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
print(f"上传文件夹路径: {UPLOAD_FOLDER}")

@app.route('/static/<path:file_name>')
def get_static_file(file_name):
    """提供静态文件"""
    static_dirs = [
        os.path.join(os.path.dirname(__file__), 'static'),
        os.path.join(os.path.dirname(__file__), '../../common/static'),
    ]

    for static_dir in static_dirs:
        file_path = os.path.join(static_dir, file_name)
        if os.path.exists(file_path):
            logger.debug(f"提供静态文件: {file_name} 从 {static_dir}")
            return send_from_directory(static_dir, file_name)

    logger.error(f"静态文件未找到: {file_name}")
    abort(404)


@app.route('/webfonts/<path:file_name>')
def get_webfonts_file(file_name):
    """提供字体文件"""
    font_file_name = f"webfonts/{file_name}"
    return get_static_file(font_file_name)



@app.route('/')
def index():
    """渲染主页面"""
    logger.info("访问首页")
    
    # 检查OnlyOffice服务是否可用
    onlyoffice_status = 'unknown'
    try:
        import requests
        # 简单检查OnlyOffice API是否可达
        test_url = f"{ONLY_OFFICE_API}/web-apps/apps/api/documents/api.js"
        response = requests.get(test_url, timeout=2)
        onlyoffice_status = 'available' if response.status_code == 200 else 'unavailable'
    except:
        onlyoffice_status = 'unavailable'
    
    return render_template(
        'index.html',
        config={
            'only_office_api': ONLY_OFFICE_API,
            'onlyoffice_status': onlyoffice_status
        }
    )


@app.route('/health')
def health_check():
    """健康检查端点"""
    try:
        # 检查数据库连接
        db = get_db()
        cursor = db.cursor()
        cursor.execute('SELECT 1')
        db_status = 'healthy'
    except Exception as e:
        logger.error(f"数据库健康检查失败: {e}")
        db_status = 'unhealthy'
    
    # 检查上传文件夹
    upload_folder_status = 'healthy' if os.path.exists(UPLOAD_FOLDER) else 'unhealthy'
    
    # 检查文档文件夹
    doc_folder_status = 'healthy' if os.path.exists(DOCUMENT_UPLOAD_FOLDER) else 'unhealthy'
    
    return jsonify({
        'status': 'healthy' if all([
            db_status == 'healthy',
            upload_folder_status == 'healthy',
            doc_folder_status == 'healthy'
        ]) else 'degraded',
        'components': {
            'database': db_status,
            'upload_folder': upload_folder_status,
            'document_folder': doc_folder_status,
            'onlyoffice': check_onlyoffice_status()
        },
        'timestamp': datetime.now().isoformat()
    })


def check_onlyoffice_status():
    """检查OnlyOffice服务状态"""
    try:
        import requests
        test_url = f"{ONLY_OFFICE_API}/health"
        response = requests.get(test_url, timeout=2)
        return 'available' if response.status_code == 200 else 'unavailable'
    except:
        return 'unavailable'


@app.route('/api/documents/upload', methods=['POST'])
def upload_document():
    """上传文档并准备预览"""
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': get_msg('online_office.no_file')}), 400

        file = request.files['file']

        if file.filename == '':
            return jsonify({'success': False, 'error': get_msg('online_office.no_file')}), 400

        # 检查文件扩展名
        file_ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else ''
        if file_ext not in ALLOWED_DOC_EXTENSIONS:
            return jsonify({
                'success': False,
                'error': get_msg('online_office.unsupported_file_type', types=", ".join(ALLOWED_DOC_EXTENSIONS))
            }), 400

        # 检查文件大小
        file.seek(0, os.SEEK_END)
        file_size = file.tell()
        file.seek(0)

        if file_size > DOC_MAX_FILE_SIZE:
            return jsonify({
                'success': False,
                'error': get_msg('online_office.file_too_large_max', max=DOC_MAX_FILE_SIZE // 1024 // 1024)
            }), 400

        # 生成唯一ID和文件名
        doc_id = str(uuid.uuid4())
        safe_filename = f"{doc_id}.{file_ext}"
        upload_path = os.path.join(DOCUMENT_UPLOAD_FOLDER, safe_filename)

        # 保存文件
        file.save(upload_path)
        logger.info(f"文档上传成功: {safe_filename}, 大小: {file_size}字节")

        # 获取文档信息
        original_filename = file.filename
        file_type = get_file_type(file_ext)

        # 获取Docker可访问的主机地址
        docker_host = get_docker_host()
        logger.info(f"使用Docker主机地址: {docker_host}")
        file_url = f"http://{docker_host}:19000/api/documents/download/{doc_id}"
        logger.info(f"文档URL（Docker）: {file_url}")

        # 存储文档信息到数据库
        timestamp = int(datetime.now().timestamp())
        upload_time = datetime.now().isoformat()
        
        db = get_db()
        cursor = db.cursor()
        cursor.execute('''
            INSERT INTO documents (id, original_filename, filename, file_type, file_ext, 
                                   size, upload_time, url, path, key, last_modified)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (doc_id, original_filename, safe_filename, file_type, file_ext, 
              file_size, upload_time, file_url, upload_path, f"{doc_id}_{timestamp}", upload_time))
        db.commit()
        
        # 获取完整的文档信息
        cursor.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        document_row = cursor.fetchone()
        document_info = dict(document_row)
        
        onlyoffice_config = generate_onlyoffice_config(document_info)
        logger.info(f"生成的OnlyOffice配置: {json.dumps(onlyoffice_config, indent=2)}")

        return jsonify({
            'success': True,
            'document': {
                'id': doc_id,
                'original_filename': original_filename,
                'file_type': file_type,
                'size': file_size,
                'url': file_url,
                'key': document_info['key'],
                'file_ext': file_ext
            },
            'onlyoffice_config': onlyoffice_config  # 返回完整配置
        })

    except Exception as e:
        logger.error(f"文档上传失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/documents/download/<doc_id>')
def download_document(doc_id):
    """提供文档下载（OnlyOffice会访问这个URL获取文档）"""
    try:
        db = get_db()
        cursor = db.cursor()
        cursor.execute('SELECT * FROM documents WHERE id = ? AND status = "active"', (doc_id,))
        doc_row = cursor.fetchone()
        
        if not doc_row:
            abort(404, get_msg('online_office.doc_not_found'))
            
        doc_info = dict(doc_row)
        file_path = doc_info['path']

        if not os.path.exists(file_path):
            abort(404, get_msg('online_office.file_not_found'))

        # 设置正确的Content-Type
        content_type = get_content_type(doc_info['file_ext'])

        return send_from_directory(
            os.path.dirname(file_path),
            os.path.basename(file_path),
            as_attachment=False,
            mimetype=content_type
        )

    except Exception as e:
        logger.error(f"下载文档失败: {str(e)}")
        abort(500)


@app.route('/api/documents/list')
def list_documents():
    """获取文档列表"""
    db = get_db()
    cursor = db.cursor()
    cursor.execute('''
        SELECT id, original_filename, file_type, file_ext, size, 
               upload_time, url, last_modified, status
        FROM documents 
        WHERE status = "active"
        ORDER BY upload_time DESC
    ''')
    
    docs = []
    for row in cursor.fetchall():
        doc = dict(row)
        # 转换时间为友好格式
        try:
            upload_time = datetime.fromisoformat(doc['upload_time'])
            doc['upload_time_formatted'] = upload_time.strftime('%Y-%m-%d %H:%M')
        except:
            doc['upload_time_formatted'] = doc['upload_time']
        
        docs.append(doc)
    
    return jsonify({'success': True, 'documents': docs})


@app.route('/api/documents/<doc_id>')
def get_document_info(doc_id):
    """获取单个文档信息"""
    db = get_db()
    cursor = db.cursor()
    cursor.execute('''
        SELECT id, original_filename, file_type, file_ext, size, 
               upload_time, url, key, last_modified, status
        FROM documents 
        WHERE id = ? AND status = "active"
    ''', (doc_id,))
    
    doc_row = cursor.fetchone()
    if not doc_row:
        return jsonify({'success': False, 'error': get_msg('online_office.doc_not_found')}), 404
    
    doc_info = dict(doc_row)
    return jsonify({'success': True, 'document': doc_info})


@app.route('/api/documents/analyze', methods=['POST'])
def analyze_document():
    """AI分析文档内容"""
    try:
        data = request.get_json()
        doc_id = data.get('doc_id')
        
        if not doc_id:
            return jsonify({'success': False, 'error': get_msg('online_office.missing_doc_id')}), 400
        
        if doc_id not in documents_db:
            return jsonify({'success': False, 'error': get_msg('online_office.doc_not_found')}), 404
        
        doc_info = documents_db[doc_id]
        file_path = doc_info['path']
        
        # 检查是否已经分析过
        db = get_db()
        cursor = db.cursor()
        cursor.execute('SELECT COUNT(*) FROM document_analysis WHERE doc_id = ?', (doc_id,))
        existing_count = cursor.fetchone()[0]
        
        if existing_count > 0:
            # 返回已保存的分析结果
            cursor.execute('SELECT * FROM document_analysis WHERE doc_id = ? ORDER BY created_time DESC', (doc_id,))
            suggestions = []
            for row in cursor.fetchall():
                suggestion = dict(row)
                suggestions.append({
                    'id': suggestion['id'],
                    'category': suggestion['category'],
                    'severity': suggestion['severity'],
                    'position': suggestion['position'],
                    'description': suggestion['description'],
                    'suggestion': suggestion['suggestion'],
                    'status': suggestion['status']
                })
        else:
            # 根据文件类型使用不同分析器
            if doc_info['file_ext'] in ['docx', 'doc']:
                # Word文档分析
                suggestions = analyze_word_document(file_path, doc_id)
            elif doc_info['file_ext'] == 'pdf':
                # PDF文档分析
                suggestions = analyze_pdf_document(file_path, doc_id)
            elif doc_info['file_ext'] in ['xlsx']:
                # Excel文档分析
                suggestions = analyze_excel_document(file_path, doc_id)
            elif doc_info['file_ext'] in ['pptx']:
                # PowerPoint文档分析
                suggestions = analyze_ppt_document(file_path, doc_id)
            elif doc_info['file_ext'] in ['txt']:
                # 文本文件分析
                suggestions = analyze_text_document(file_path, doc_id)
            else:
                suggestions = []
        
        return jsonify({
            'success': True,
            'suggestions': suggestions,
            'document': {
                'id': doc_id,
                'filename': doc_info['original_filename'],
                'type': doc_info['file_type']
            }
        })
        
    except Exception as e:
        logger.error(f"文档分析失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


def analyze_word_document(file_path, doc_id):
    """分析Word文档并保存结果到数据库"""
    try:
        import docx
        from docx import Document
        
        doc = Document(file_path)
        suggestions = []
        analysis_time = datetime.now().isoformat()
        
        # 检查文档结构
        if len(doc.paragraphs) < 2:
            suggestions.append({
                'category': '结构',
                'severity': '中',
                'position': '文档结构',
                'description': '文档内容较少，建议增加更多段落',
                'suggestion': '添加介绍、正文和结论部分'
            })
        
        # 检查标题
        has_title = False
        for para in doc.paragraphs[:5]:
            if para.style.name.startswith('Heading'):
                has_title = True
                break
        
        if not has_title:
            suggestions.append({
                'category': '格式',
                'severity': '低',
                'position': '文档开头',
                'description': '文档缺少标题',
                'suggestion': '使用"标题1"样式添加文档标题'
            })
        
        # 检查段落长度
        for i, para in enumerate(doc.paragraphs):
            if para.style.name == 'Normal' and len(para.text.strip()) > 500:
                suggestions.append({
                    'category': '可读性',
                    'severity': '低',
                    'position': f'第{i+1}段',
                    'description': '段落过长，影响可读性',
                    'suggestion': '将长段落拆分为2-3个段落'
                })
        
        # 检查拼写错误（简单示例）
        common_mistakes = {
            'teh': 'the',
            'adn': 'and',
            'thier': 'their',
            'recieve': 'receive',
            'seperate': 'separate'
        }
        
        for para in doc.paragraphs:
            text = para.text.lower()
            for mistake, correction in common_mistakes.items():
                if mistake in text:
                    suggestions.append({
                        'category': '语法',
                        'severity': '低',
                        'position': '文档内容',
                        'description': f'常见拼写错误: "{mistake}"',
                        'suggestion': f'建议改为: "{correction}"'
                    })
                    break
        
        # 保存分析结果到数据库
        db = get_db()
        cursor = db.cursor()
        
        for suggestion in suggestions[:10]:  # 最多保存10条建议
            cursor.execute('''
                INSERT INTO document_analysis (doc_id, category, severity, position, 
                                               description, suggestion, created_time)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (doc_id, suggestion['category'], suggestion['severity'], 
                  suggestion['position'], suggestion['description'], 
                  suggestion['suggestion'], analysis_time))
        
        # 更新文档的分析时间
        cursor.execute('''
            UPDATE documents 
            SET analysis_time = ?
            WHERE id = ?
        ''', (analysis_time, doc_id))
        
        db.commit()
        
        # 返回带ID的建议
        cursor.execute('SELECT * FROM document_analysis WHERE doc_id = ? ORDER BY created_time DESC', (doc_id,))
        saved_suggestions = []
        for row in cursor.fetchall():
            suggestion = dict(row)
            saved_suggestions.append({
                'id': suggestion['id'],
                'category': suggestion['category'],
                'severity': suggestion['severity'],
                'position': suggestion['position'],
                'description': suggestion['description'],
                'suggestion': suggestion['suggestion']
            })
        
        return saved_suggestions
        
    except Exception as e:
        logger.error(f"Word文档分析失败: {str(e)}")
        # 保存错误信息
        error_suggestion = {
            'category': '系统',
            'severity': '低',
            'position': '文档分析',
            'description': f'文档分析遇到问题: {str(e)}',
            'suggestion': '请确保文档格式正确且可访问'
        }
        
        try:
            db = get_db()
            cursor = db.cursor()
            cursor.execute('''
                INSERT INTO document_analysis (doc_id, category, severity, position, 
                                               description, suggestion, created_time)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (doc_id, error_suggestion['category'], error_suggestion['severity'], 
                  error_suggestion['position'], error_suggestion['description'], 
                  error_suggestion['suggestion'], datetime.now().isoformat()))
            db.commit()
        except:
            pass
            
        return [error_suggestion]


def analyze_pdf_document(file_path):
    """分析PDF文档"""
    try:
        # 这里可以集成PyPDF2等库进行PDF分析
        # 目前返回模拟数据
        return [{
            'id': 1,
            'category': 'PDF',
            'severity': '低',
            'position': '文档分析',
            'description': 'PDF文档分析需要额外依赖库',
            'suggestion': '建议安装PyPDF2或pdfplumber库进行深度分析'
        }]
    except Exception as e:
        logger.error(f"PDF文档分析失败: {str(e)}")
        return []


def analyze_excel_document(file_path):
    """分析Excel文档"""
    try:
        import pandas as pd
        
        # 尝试读取Excel
        xls = pd.ExcelFile(file_path)
        suggestions = []
        
        # 检查工作表数量
        if len(xls.sheet_names) == 0:
            suggestions.append({
                'id': 1,
                'category': '结构',
                'severity': '高',
                'position': '工作簿',
                'description': 'Excel文件为空',
                'suggestion': '添加至少一个工作表'
            })
        else:
            suggestions.append({
                'id': 1,
                'category': '信息',
                'severity': '低',
                'position': '工作簿',
                'description': f'包含 {len(xls.sheet_names)} 个工作表',
                'suggestion': '工作表结构正常'
            })
        
        return suggestions
        
    except Exception as e:
        logger.error(f"Excel文档分析失败: {str(e)}")
        return []


def analyze_ppt_document(file_path):
    """分析PowerPoint文档"""
    return [{
        'id': 1,
        'category': 'PPT',
        'severity': '低',
        'position': '文档分析',
        'description': 'PowerPoint文档分析',
        'suggestion': '建议添加更多幻灯片以提高演示效果'
    }]


def analyze_text_document(file_path, doc_id):
    """分析文本文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        suggestions = []
        analysis_time = datetime.now().isoformat()
        
        # 检查文件大小
        if len(content) < 100:
            suggestions.append({
                'category': '内容',
                'severity': '中',
                'position': '文档整体',
                'description': '文本内容较少',
                'suggestion': '建议增加更多内容'
            })
        
        # 检查段落分割
        paragraphs = content.split('\n\n')
        if len(paragraphs) == 1 and len(content) > 500:
            suggestions.append({
                'category': '格式',
                'severity': '低',
                'position': '文档格式',
                'description': '缺少段落分割',
                'suggestion': '使用空行分割段落提高可读性'
            })
        
        # 保存到数据库
        db = get_db()
        cursor = db.cursor()
        
        for suggestion in suggestions:
            cursor.execute('''
                INSERT INTO document_analysis (doc_id, category, severity, position, 
                                               description, suggestion, created_time)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (doc_id, suggestion['category'], suggestion['severity'], 
                  suggestion['position'], suggestion['description'], 
                  suggestion['suggestion'], analysis_time))
        
        cursor.execute('''
            UPDATE documents 
            SET analysis_time = ?
            WHERE id = ?
        ''', (analysis_time, doc_id))
        
        db.commit()
        
        # 返回结果
        cursor.execute('SELECT * FROM document_analysis WHERE doc_id = ? ORDER BY created_time DESC', (doc_id,))
        saved_suggestions = []
        for row in cursor.fetchall():
            suggestion = dict(row)
            saved_suggestions.append({
                'id': suggestion['id'],
                'category': suggestion['category'],
                'severity': suggestion['severity'],
                'position': suggestion['position'],
                'description': suggestion['description'],
                'suggestion': suggestion['suggestion']
            })
        
        return saved_suggestions
        
    except Exception as e:
        logger.error(f"文本文件分析失败: {str(e)}")
        return []


# 文档管理API
@app.route('/api/documents/delete/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    """删除文档（软删除）"""
    try:
        db = get_db()
        cursor = db.cursor()
        
        # 检查文档是否存在
        cursor.execute('SELECT * FROM documents WHERE id = ?', (doc_id,))
        doc_row = cursor.fetchone()
        
        if not doc_row:
            return jsonify({'success': False, 'error': get_msg('online_office.doc_not_found')}), 404
        
        # 软删除：标记为已删除
        cursor.execute('''
            UPDATE documents 
            SET status = 'deleted', last_modified = ?
            WHERE id = ?
        ''', (datetime.now().isoformat(), doc_id))
        
        db.commit()
        
        return jsonify({'success': True, 'message': get_msg('online_office.doc_deleted')})
        
    except Exception as e:
        logger.error(f"删除文档失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/documents/update-suggestion', methods=['POST'])
def update_suggestion_status():
    """更新建议状态（接受/忽略）"""
    try:
        data = request.get_json()
        suggestion_id = data.get('suggestion_id')
        status = data.get('status')  # accepted, ignored
        
        if not suggestion_id or status not in ['accepted', 'ignored']:
            return jsonify({'success': False, 'error': get_msg('online_office.param_error')}), 400
        
        db = get_db()
        cursor = db.cursor()
        
        # 更新建议状态
        cursor.execute('''
            UPDATE document_analysis 
            SET status = ?
            WHERE id = ?
        ''', (status, suggestion_id))
        
        db.commit()
        
        return jsonify({'success': True, 'message': get_msg('online_office.suggestion_status_updated', status=status)})
        
    except Exception as e:
        logger.error(f"更新建议状态失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/documents/stats')
def get_document_stats():
    """获取文档统计信息"""
    try:
        db = get_db()
        cursor = db.cursor()
        
        # 获取文档总数
        cursor.execute('SELECT COUNT(*) FROM documents WHERE status = "active"')
        total_docs = cursor.fetchone()[0]
        
        # 获取已分析文档数
        cursor.execute('SELECT COUNT(DISTINCT doc_id) FROM document_analysis')
        analyzed_docs = cursor.fetchone()[0]
        
        # 获取建议统计
        cursor.execute('''
            SELECT severity, COUNT(*) as count 
            FROM document_analysis 
            GROUP BY severity
        ''')
        severity_stats = {}
        for row in cursor.fetchall():
            severity_stats[row['severity']] = row['count']
        
        # 获取文件类型统计
        cursor.execute('''
            SELECT file_ext, COUNT(*) as count 
            FROM documents 
            WHERE status = "active"
            GROUP BY file_ext
        ''')
        file_type_stats = {}
        for row in cursor.fetchall():
            file_type_stats[row['file_ext']] = row['count']
        
        return jsonify({
            'success': True,
            'stats': {
                'total_documents': total_docs,
                'analyzed_documents': analyzed_docs,
                'severity_distribution': severity_stats,
                'file_type_distribution': file_type_stats
            }
        })
        
    except Exception as e:
        logger.error(f"获取统计信息失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/debug/jwt', methods=['GET'])
def debug_jwt():
    """调试JWT生成"""
    test_payload = {
        "document": {
            "fileType": "docx",
            "key": "test_key",
            "title": "test.docx",
            "url": "http://localhost:19000/api/documents/download/test"
        }
    }

    token = generate_jwt_token(test_payload)

    # 验证令牌
    try:
        decoded = jwt.decode(token, JWT_SECRET, algorithms=['HS256'])
        return jsonify({
            'success': True,
            'token': token,
            'decoded': decoded,
            'secret_used': JWT_SECRET
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'token': token
        })


@app.route('/callback', methods=['POST'])
def onlyoffice_callback():
    """处理OnlyOffice的回调请求"""
    try:
        data = request.json
        logger.info(f"收到OnlyOffice回调: {data}")

        # 获取回调类型
        status = data.get('status', 0)

        if status == 0:  # 用户正在查看文档
            logger.info("用户正在查看文档")
            return jsonify({"error": 0})

        elif status == 1:  # 文档已编辑
            logger.info("文档已被编辑")
            # 可以在这里处理编辑事件
            return jsonify({"error": 0})

        elif status == 2:  # 文档已保存（重要！）
            logger.info("文档已保存")

            # 获取文档URL
            if 'url' in data:
                download_url = data['url']
                logger.info(f"文档下载URL: {download_url}")

                # 这里可以下载文档到服务器
                # 或者记录保存事件

            return jsonify({"error": 0})

        elif status == 3:  # 保存文档时出错
            logger.error("保存文档时出错")
            return jsonify({"error": 0})

        elif status == 4:  # 文档关闭且未保存
            logger.info("文档关闭且未保存")
            return jsonify({"error": 0})

        elif status == 6:  # 用户正在编辑文档
            users = data.get('users', [])
            logger.info(f"用户正在编辑文档: {users}")
            return jsonify({"error": 0})

        elif status == 7:  # 强制保存请求
            logger.info("收到强制保存请求")
            return jsonify({"error": 0})

        else:
            logger.warning(f"未知的回调状态: {status}")
            return jsonify({"error": 0})

    except Exception as e:
        logger.error(f"处理回调失败: {str(e)}")
        return jsonify({"error": 1, "message": str(e)})


@app.route('/api/documents/save/<doc_id>', methods=['POST'])
def save_document(doc_id):
    """保存文档（从OnlyOffice下载）"""
    try:
        if doc_id not in documents_db:
            return jsonify({'success': False, 'error': get_msg('online_office.doc_not_found')}), 404

        # 从请求中获取文档数据
        file_data = request.data

        if not file_data:
            return jsonify({'success': False, 'error': get_msg('online_office.no_document_data')}), 400

        # 保存文档到服务器
        doc_info = documents_db[doc_id]
        file_path = doc_info['path']

        # 备份原文件（可选）
        backup_path = file_path + '.backup'
        if os.path.exists(file_path):
            shutil.copy2(file_path, backup_path)

        # 保存新文件
        with open(file_path, 'wb') as f:
            f.write(file_data)

        logger.info(f"文档已保存: {doc_info['original_filename']}")

        return jsonify({
            'success': True,
            'message': get_msg('online_office.doc_saved')
        })

    except Exception as e:
        logger.error(f"保存文档失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


if __name__ == '__main__':

    JWT_AVAILABLE = True
    logger.info("启动Flask应用...")
    app.run(
        debug=False,
        host='0.0.0.0',
        port=19000,
        threaded=True
    )