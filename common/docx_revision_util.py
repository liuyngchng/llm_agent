#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word docx 修订模式（Track Changes）工具模块。

基于 python-docx + lxml 直接操作 OOXML 修订标记（w:ins / w:del），
支持以修订模式对 Word 文档进行插入、删除、替换操作，
以及接受/拒绝全部修订。

用法示例:
    from common.docx_revision_util import (
        tracked_insert_text, tracked_delete_text, tracked_replace_text,
        tracked_replace_in_document, accept_all_changes, reject_all_changes,
        get_tracked_changes_summary,
    )

    doc = Document('report.docx')
    para = doc.paragraphs[0]
    tracked_insert_text(doc, para, '新增内容', author='Reviewer')
    tracked_replace_text(doc, para, '旧文本', '新文本', author='Reviewer')
    doc.save('report_revised.docx')
"""

import datetime
from lxml import etree
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx import Document


# ── internal helpers ──────────────────────────────────────────────

def _next_rev_id(doc: Document) -> int:
    """扫描文档中已有的修订 ID，返回下一个可用 ID。"""
    body = doc.element.body
    ids = body.xpath('.//w:ins/@w:id | .//w:del/@w:id')
    if ids:
        return max(int(i) for i in ids) + 1
    return 1


def _make_ins(doc: Document, author: str) -> etree.Element:
    """创建 w:ins（插入修订）标记。"""
    rev_id = _next_rev_id(doc)
    date_str = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'), str(rev_id))
    ins.set(qn('w:author'), author)
    ins.set(qn('w:date'), date_str)
    return ins


def _make_del(doc: Document, author: str) -> etree.Element:
    """创建 w:del（删除修订）标记。"""
    rev_id = _next_rev_id(doc)
    date_str = datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    del_elem = OxmlElement('w:del')
    del_elem.set(qn('w:id'), str(rev_id))
    del_elem.set(qn('w:author'), author)
    del_elem.set(qn('w:date'), date_str)
    return del_elem


def _make_run(text: str, bold: bool = False, italic: bool = False,
              underline: bool = False) -> etree.Element:
    """创建 w:r（文本运行）元素，可选加粗/斜体/下划线格式。"""
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    has_fmt = False
    if bold:
        rPr.append(OxmlElement('w:b'))
        has_fmt = True
    if italic:
        rPr.append(OxmlElement('w:i'))
        has_fmt = True
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
        has_fmt = True
    if has_fmt:
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r


# ── public API ────────────────────────────────────────────────────

def tracked_insert_text(doc: Document, paragraph, text: str,
                         author: str = "DocForge", bold: bool = False):
    """
    在段落末尾以修订模式（Track Changes）插入文本。

    Args:
        doc: python-docx Document 对象
        paragraph: 目标 Paragraph 对象
        text: 要插入的文本
        author: 修订者名称
        bold: 是否加粗

    Example:
        >>> doc = Document('report.docx')
        >>> tracked_insert_text(doc, doc.paragraphs[0], '新增段落内容')
        >>> doc.save('report_revised.docx')
    """
    ins = _make_ins(doc, author)
    run = _make_run(text, bold=bold)
    ins.append(run)
    paragraph._element.append(ins)


def tracked_delete_paragraph(doc: Document, paragraph, author: str = "DocForge"):
    """
    以修订模式删除整个段落。

    将段落中所有文本内容包裹在 w:del 标记中，Word 中显示为删除。

    Args:
        doc: python-docx Document 对象
        paragraph: 要删除的 Paragraph 对象
        author: 修订者名称
    """
    p_elem = paragraph._element
    del_elem = _make_del(doc, author)
    children = list(p_elem)
    for child in children:
        p_elem.remove(child)
        del_elem.append(child)
    p_elem.append(del_elem)


def tracked_delete_text(doc: Document, paragraph, text: str,
                         author: str = "DocForge") -> bool:
    """
    以修订模式删除段落中匹配的文本。

    支持跨 run 的文本匹配（python-docx 中段落文本可能被拆分在多个 run 中）。

    Args:
        doc: python-docx Document 对象
        paragraph: 目标 Paragraph 对象
        text: 要删除的文本（精确匹配）
        author: 修订者名称

    Returns:
        bool: 是否找到并标记删除
    """
    full_text = paragraph.text
    if text not in full_text:
        return False

    start_idx = full_text.index(text)
    end_idx = start_idx + len(text)

    p_elem = paragraph._element
    runs = p_elem.xpath('w:r')

    # 收集每个 run 在段落全文中的起止位置
    run_info = []
    pos = 0
    for r in runs:
        t_elem = r.find(qn('w:t'))
        t_text = t_elem.text if t_elem is not None and t_elem.text else ''
        run_info.append({'elem': r, 'text': t_text, 'start': pos, 'end': pos + len(t_text)})
        pos += len(t_text)

    # 找到与删除区间有交集的 run
    affected = [ri for ri in run_info if ri['end'] > start_idx and ri['start'] < end_idx]
    if not affected:
        return False

    for ri in affected:
        run_start = max(ri['start'], start_idx) - ri['start']
        run_end = min(ri['end'], end_idx) - ri['start']

        if run_start == 0 and run_end >= len(ri['text']):
            # 整个 run 被覆盖 — 直接包裹
            del_elem = _make_del(doc, author)
            p_elem.remove(ri['elem'])
            del_elem.append(ri['elem'])
            p_elem.append(del_elem)
        else:
            # 部分覆盖 — 拆分为 before / deleted / after
            original = ri['text']
            before = original[:run_start]
            deleted = original[run_start:run_end]
            after = original[run_end:]

            new_elems = []
            if before:
                new_elems.append(_make_run(before))
            if deleted:
                del_elem = _make_del(doc, author)
                del_elem.append(_make_run(deleted))
                new_elems.append(del_elem)
            if after:
                new_elems.append(_make_run(after))

            parent = p_elem
            idx = list(parent).index(ri['elem'])
            parent.remove(ri['elem'])
            for i, elem in enumerate(new_elems):
                parent.insert(idx + i, elem)

    return True


def tracked_replace_text(doc: Document, paragraph, old_text: str, new_text: str,
                          author: str = "DocForge") -> bool:
    """
    以修订模式替换段落中的文本（标记删除旧文本 + 插入新文本）。

    新文本插入到被删除文本的原位置之后。

    Args:
        doc: python-docx Document 对象
        paragraph: 目标 Paragraph 对象
        old_text: 要替换的原文本
        new_text: 替换后的新文本
        author: 修订者名称

    Returns:
        bool: 是否成功替换

    Example:
        >>> doc = Document('report.docx')
        >>> tracked_replace_text(doc, doc.paragraphs[2], '2024', '2025', author='Editor')
        >>> doc.save('report_revised.docx')
    """
    if old_text not in paragraph.text:
        return False

    deleted = tracked_delete_text(doc, paragraph, old_text, author)
    if not deleted:
        return False

    # 在最后一个 w:del 之后插入新文本，保持位置接近
    ins = _make_ins(doc, author)
    ins.append(_make_run(new_text))
    del_elems = paragraph._element.xpath('w:del')
    if del_elems:
        last_del = del_elems[-1]
        idx = list(paragraph._element).index(last_del) + 1
        paragraph._element.insert(idx, ins)
    else:
        paragraph._element.append(ins)

    return True


def tracked_replace_in_document(doc: Document, old_text: str, new_text: str,
                                 author: str = "DocForge", max_replace: int = 0) -> int:
    """
    在整个文档中以修订模式替换所有匹配的文本。

    遍历所有段落（含表格），查找并替换每一个匹配项。

    Args:
        doc: python-docx Document 对象
        old_text: 要查找并替换的原文本
        new_text: 替换后的新文本
        author: 修订者名称
        max_replace: 最大替换次数，0 表示不限制

    Returns:
        int: 实际替换的处数

    Example:
        >>> doc = Document('report.docx')
        >>> count = tracked_replace_in_document(doc, '旧公司名', '新公司名', author='Editor')
        >>> print(f'共替换 {count} 处')
        >>> doc.save('report_revised.docx')
    """
    count = 0
    safe_limit = 1000

    for paragraph in doc.paragraphs:
        iterations = 0
        while old_text in paragraph.text and iterations < safe_limit:
            iterations += 1
            if tracked_replace_text(doc, paragraph, old_text, new_text, author):
                count += 1
                if max_replace > 0 and count >= max_replace:
                    return count
            else:
                break

    # 同时处理表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    iterations = 0
                    while old_text in paragraph.text and iterations < safe_limit:
                        iterations += 1
                        if tracked_replace_text(doc, paragraph, old_text, new_text, author):
                            count += 1
                            if max_replace > 0 and count >= max_replace:
                                return count
                        else:
                            break

    return count


def accept_all_changes(doc: Document):
    """
    接受文档中所有修订。

    - 接受插入：将 w:ins 内的内容展开为普通文本
    - 接受删除：移除 w:del 及其内容

    Args:
        doc: python-docx Document 对象
    """
    body = doc.element.body

    for ins in body.xpath('.//w:ins'):
        parent = ins.getparent()
        if parent is not None:
            idx = list(parent).index(ins)
            children = list(ins)
            for i, child in enumerate(children):
                parent.insert(idx + i, child)
            parent.remove(ins)

    for del_elem in body.xpath('.//w:del'):
        parent = del_elem.getparent()
        if parent is not None:
            parent.remove(del_elem)


def reject_all_changes(doc: Document):
    """
    拒绝文档中所有修订（恢复到修订前状态）。

    - 拒绝插入：移除 w:ins 及其内容
    - 拒绝删除：将 w:del 内的内容展开为普通文本

    Args:
        doc: python-docx Document 对象
    """
    body = doc.element.body

    for ins in body.xpath('.//w:ins'):
        parent = ins.getparent()
        if parent is not None:
            parent.remove(ins)

    for del_elem in body.xpath('.//w:del'):
        parent = del_elem.getparent()
        if parent is not None:
            idx = list(parent).index(del_elem)
            children = list(del_elem)
            for i, child in enumerate(children):
                parent.insert(idx + i, child)
            parent.remove(del_elem)


def get_tracked_changes_summary(doc: Document) -> dict:
    """
    获取文档中修订的摘要信息。

    Returns:
        dict: {'insertions': 插入处数, 'deletions': 删除处数, 'authors': [修订者列表]}

    Example:
        >>> summary = get_tracked_changes_summary(doc)
        >>> print(f"插入 {summary['insertions']} 处，删除 {summary['deletions']} 处")
    """
    body = doc.element.body
    insertions = len(body.xpath('.//w:ins'))
    deletions = len(body.xpath('.//w:del'))

    authors = set()
    for elem in body.xpath('.//w:ins') + body.xpath('.//w:del'):
        author = elem.get(qn('w:author'))
        if author:
            authors.add(author)

    return {
        'insertions': insertions,
        'deletions': deletions,
        'authors': sorted(authors),
    }
