#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.

"""
pip install python-docx
处理docx的段落、标题、标题下的文本、文档目录、
修改docx的文档内容、设定新增的docx文档的文本格式等
"""
import logging.config
import os
import re
import time

from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH

from common.const import MIN_DESC_TXT_LEN
from common.vdb_util import search_txt
from common.sys_init import init_yml_cfg

log_config_path = 'logging.conf'
if os.path.exists(log_config_path):
    logging.config.fileConfig(log_config_path, encoding="utf-8")
else:
    from common.const import LOG_FORMATTER
    logging.basicConfig(level=logging.INFO,format= LOG_FORMATTER, force=True)
logger = logging.getLogger(__name__)


def get_reference_from_vdb(keywords: str, vdb_dir: str, sys_cfg: dict) -> str:
    """
    获取vdb中与关键词相关的文本
    :param keywords: 关键词
    :param vdb_dir: 向量数据库目录
    :param sys_cfg: 系统配置
    :return: 文本
    """
    logger.debug(f"vdb_dir, {vdb_dir}")
    reference = ""
    if not vdb_dir:
        return reference

    try:
        if "" != vdb_dir and os.path.exists(vdb_dir):
            reference = search_txt(keywords, vdb_dir, 0.2, sys_cfg, 2).strip()
        else:
            logger.warning(f"vdb_dir_not_exist: {vdb_dir}, get no references")
            reference = ""
        # logging.info(f"vdb_get_txt:\n{reference}\nby_search_{keywords}")
    except Exception as exp:
        logger.exception(f"get_references_from_vdb_failed, {keywords}")
    return reference


def extract_catalogue(target_doc: str) -> str:
    """
    生成 docx 文档的三级目录清单
    :param target_doc: 目标文档
    :return: 目录文本
    """
    doc = Document(target_doc)
    catalogue_lines = []
    # 初始化各级标题计数器 [一级, 二级, 三级]
    level_counters = [0, 0, 0]
    for para in doc.paragraphs:
        style_name = para.style.name.lower()  # 统一转换为小写
        # 检查是否为标题（兼容中英文样式名）
        if not style_name.startswith(('heading', '标题')):
            continue
        # 提取标题级别数字
        level_str = ''.join(filter(str.isdigit, style_name))
        if not level_str:
            continue
        level = int(level_str)
        if level <1 or level > 3:
            continue
        # 仅处理1-3级标题
        # 更新计数器：当前级别+1，更低级别清零
        level_index = level - 1
        level_counters[level_index] += 1
        for i in range(level_index + 1, 3):
            level_counters[i] = 0
        # 生成编号 (如 "2.1.1")
        number_parts = []
        for i in range(level):
            number_parts.append(str(level_counters[i]))
        number_str = '.'.join(number_parts)
        # 添加缩进和目录行
        indent = "  " * (level - 1)
        catalogue_lines.append(f"{indent}{number_str} {para.text}")
    return "\n".join(catalogue_lines)

def refresh_current_heading(para: Paragraph, heading: list) -> bool:
    """
    更新当前层级标题，并返回当前段落内容是否为文章标题
    :param para: 段落
    :param heading: 当前标题
    :return: 是否为三级标题
    """
    if "Heading" not in para.style.name:
        return False

    level = int(para.style.name.split()[-1])  # 提取数字
    headings = {1: [], 2: [], 3: [], 4: [], 5: []}  # 按需扩展层级
    headings[level].append(para.text)
    logger.debug(f"heading_part_caught: H{level}: {para.text}")
    if None == heading:
        logger.error("heading_list_cant't_be_refresh")
        return True
    if len(heading) != 0:
        heading.pop()
    heading.append(para.text)
    return True

def is_3rd_heading(para: Paragraph) -> bool:
    """
    判断是否为三级标题
    :param para: 段落
    :return: 是否为三级标题
    """
    if "Heading" not in para.style.name:
        return False

    level = int(para.style.name.split()[-1])  # 提取数字
    if level == 3:
        return True
    else:
        return False

def get_outline_txt(file_name: str) -> str:
    """
    获取word文档的三级目录，输出格式如下：
    # 1. 一级目录
    ## 1.1 二级目录
    ### 1.1.1 三级目录
    """
    doc = Document(file_name)
    result_lines = []  # 存储每行目录文本
    chapter_num = 0    # 一级标题编号
    section_num = 0    # 二级标题编号
    subsection_num = 0 # 三级标题编号

    for para in doc.paragraphs:
        style_name = para.style.name.lower()
        # 跳过非标题段落
        if not style_name.startswith(('heading', '标题')):
            continue

        # 提取标题级别
        level_str = ''.join(filter(str.isdigit, style_name))
        if not level_str:
            continue
        level = int(level_str)

        # 只处理1-3级标题
        if level < 1 or level > 3:
            continue

        text = para.text.strip()
        if not text:
            continue
        # 处理一级标题
        if level == 1:
            chapter_num += 1
            section_num = 0
            subsection_num = 0
            result_lines.append(f"# {chapter_num}. {text}")

        # 处理二级标题
        elif level == 2 and chapter_num > 0:
            section_num += 1
            subsection_num = 0
            result_lines.append(f"## {chapter_num}.{section_num} {text}")

        # 处理三级标题
        elif level == 3 and chapter_num > 0 and section_num > 0:
            subsection_num += 1
            result_lines.append(f"### {chapter_num}.{section_num}.{subsection_num} {text}")

    return "\n".join(result_lines)

def is_txt_para(para: Paragraph, current_heading:list) -> bool:
    """
    判断当前段落是否为描述性的文本
    :param para: 写作要求word文档中的一个段落
    :param current_heading: 当前para 所在的目录
    return False： 不是写作要求； True： 是写作要求
    """

    pattern = r'^(图|表)\s*\d+[\.\-\s]'  # 匹配"图1."/"表2-"等开头
    if refresh_current_heading(para, current_heading):
        return False
    if "TOC" in para.style.name or para._element.xpath(".//w:instrText[contains(.,'TOC')]"):
        logger.info(f"doc_table_of_content: {para.text}")
        return False
    if "Caption" in para.style.name or re.match(pattern, para.text):
        logger.info(f"table_or_picture_title: {para.text}")
        return False
    if not para.text:
        return False
    if len(para.text.strip()) < MIN_DESC_TXT_LEN:
        logger.info(f"ignored_short_txt {para.text}")
        return False
    if len(current_heading) == 0 or len(current_heading[0]) < 2:
        logger.info(f"heading_err_for_para, {current_heading}, {para.text}")
        return False
    return True


def gen_docx_template_with_outline_txt(task_id: int, os_dir:str, title: str, outline: str) -> str:
    """
    生成Word docx 文档文件，并返回文件名称
    :param task_id: 执行任务的ID
    :param os_dir: 输出的本地文件目录
    :param title: 文档标题
    :param outline: 目录文本信息，结构如下，至多支持5级\n1. 一级标题\n  1.1 二级标题\n    1.1.1 三级标题\n
    即outline中的各行文本中，一级标题前没有空格，二级标题前有2个空格，三级标题前有4个空格
    :return: 包含三级目录的Word docx 文档文件的绝对路径名称
    """
    doc = Document()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.size = Pt(28)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加分页符使目录从新页面开始
    doc.add_page_break()
    lines = outline.strip().split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        indent_level = 0
        for char in line:
            if char == ' ':
                indent_level += 1
            elif not char.isspace():
                break
        # 添加标题并获取段落对象
        if indent_level == 0:
            level = 1
        elif indent_level == 2:
            level = 2
        elif indent_level == 4:
            level = 3
        elif indent_level == 6:
            level = 4
        elif indent_level >= 8:
            level = 5
        else:
            level = 5
        p = doc.add_heading(stripped, level=level)
        # 统一设置格式：黑体、黑色
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 中文字体
    # 保存文件
    filename = f"{task_id}_outline.docx"
    save_path = os.path.join(os_dir, filename)
    doc.save(save_path)
    abs_path = os.path.abspath(save_path)
    return abs_path


def get_elapsed_time(start_timestamp: float) -> str:
    """
    计算任务处理时间
    :param start_timestamp: 任务开始时间戳
    """
    try:
        current_time = time.time()
        elapsed_seconds = current_time - start_timestamp/1000
        if elapsed_seconds < 0:
            return "用时计算错误"
        minutes = int(elapsed_seconds // 60)
        seconds = int(elapsed_seconds % 60)
        if minutes == 0:
            return f"用时 {seconds} 秒"
        else:
            return f"用时 {minutes} 分 {seconds} 秒"
    except Exception as e:
        logger.error(f"时间计算错误: {str(e)}")
        return "用时未知"

if __name__ == "__main__":
    my_cfg = init_yml_cfg()
    my_source_dir = "/home/rd/doc/文档生成/knowledge_base"
    # my_target_doc = "/home/rd/doc/文档生成/output_template.docx"
    my_target_doc = "/home/rd/doc/文档生成/2.docx"
    # test_catalogue = extract_catalogue(my_target_doc)
    # logger.info(f"doc_catalogue: {test_catalogue}")
    my_doc_ctx = "我正在写一个可行性研究报告"
    doc_catalogue = extract_catalogue(my_target_doc)
    logger.info(f"my_target_doc_catalogue: {doc_catalogue}")
    my_vdb_dir = "./vdb/vdb_idx_332987902_26"
    my_task_id = int(time.time())
    output_file = 'doc_output.docx'
