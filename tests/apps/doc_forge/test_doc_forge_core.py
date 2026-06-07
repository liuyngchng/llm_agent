#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Core reliability tests for doc_forge — system prompt and code execution."""

import os
import sys
import tempfile

# Ensure the repo root is on the path
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))))

from apps.doc_forge.chat_util import build_doc_processing_system_prompt
from apps.doc_forge.code_executor import execute_code, extract_python_blocks


# ── System prompt tests ──────────────────────────────────────────────

def test_system_prompt_no_revision_references():
    """The system prompt must NOT mention revision mode or tracked changes."""
    prompt = build_doc_processing_system_prompt(
        file_paths=["/tmp/test.docx"],
        output_dir="/tmp/output",
        upload_dir="/tmp/upload",
    )
    forbidden = [
        "revision", "tracked_", "Track Changes", "修订",
        "direct_tracked_replace", "docx_revision_util",
        "accept_all_changes", "reject_all_changes",
    ]
    for term in forbidden:
        assert term not in prompt, f"Forbidden term '{term}' found in system prompt"


def test_system_prompt_recommends_python_docx():
    """The system prompt must recommend python-docx for docx operations."""
    prompt = build_doc_processing_system_prompt(
        file_paths=[],
        output_dir="/tmp/output",
    )
    assert "python-docx" in prompt
    # 提示词应描述库的能力（段落、表格），而非提供示例代码
    assert "段落" in prompt


def test_system_prompt_forbids_markdown_roundtrip():
    """The system prompt must forbid markdown roundtrip."""
    prompt = build_doc_processing_system_prompt(file_paths=[], output_dir="/tmp/output")
    assert "禁止" in prompt and "markdown" in prompt.lower()


def test_system_prompt_includes_file_variables():
    """When files exist, prompt includes FILE_ variables."""
    prompt = build_doc_processing_system_prompt(
        file_paths=["/tmp/a.docx", "/tmp/b.xlsx"],
        output_dir="/tmp/output",
    )
    assert "FILE_1" in prompt
    assert "FILE_2" in prompt
    assert "/tmp/a.docx" in prompt


# ── Code executor tests ──────────────────────────────────────────────

def test_extract_python_blocks():
    """Extract Python code blocks from markdown text."""
    text = """
一些说明文字

```python
# 修改文档
from docx import Document
doc = Document("test.docx")
doc.save("output.docx")
```

更多文字
"""
    blocks = extract_python_blocks(text)
    assert len(blocks) == 1
    assert "# 修改文档" in blocks[0]
    assert "from docx import Document" in blocks[0]


def test_extract_python_blocks_multiple():
    """Extract multiple Python code blocks."""
    text = """
```python
print("block 1")
```

```python
print("block 2")
```
"""
    blocks = extract_python_blocks(text)
    assert len(blocks) == 2


def test_extract_python_blocks_empty():
    """No Python code blocks."""
    blocks = extract_python_blocks("Just some text, no code blocks here.")
    assert blocks == []


def test_execute_simple_docx_modification():
    """End-to-end: create a simple docx, modify it via execute_code, verify."""
    with tempfile.TemporaryDirectory() as tmpdir:
        # First, create a source docx
        create_code = f"""
from docx import Document
doc = Document()
doc.add_paragraph("Hello World")
doc.add_paragraph("This is a test document.")
doc.save(os.path.join(OUTPUT_DIR, "source.docx"))
print("source created")
"""
        result = execute_code(create_code, output_dir=tmpdir, upload_dir=tmpdir)
        assert result['success'], f"Source creation failed: {result['stderr']}"
        assert "source.docx" in result.get('new_files', [])

        source_path = os.path.join(tmpdir, "source.docx")
        assert os.path.exists(source_path)

        # Now modify it
        modify_code = f"""
from docx import Document
import os

src = os.path.join(OUTPUT_DIR, "source.docx")
dst = os.path.join(OUTPUT_DIR, "modified.docx")

doc = Document(src)
# Replace text in first paragraph
doc.paragraphs[0].text = "Hello Modified World"
doc.save(dst)
print("modification done")
"""
        result = execute_code(modify_code, output_dir=tmpdir, upload_dir=tmpdir)
        assert result['success'], f"Modification failed: {result['stderr']}"
        assert "modified.docx" in result.get('new_files', [])

        # Verify the modification
        from docx import Document
        modified_path = os.path.join(tmpdir, "modified.docx")
        doc = Document(modified_path)
        assert doc.paragraphs[0].text == "Hello Modified World"
        assert doc.paragraphs[1].text == "This is a test document."


def test_execute_code_with_file_paths():
    """execute_code injects FILE_ variables when file_paths provided."""
    with tempfile.TemporaryDirectory() as tmpdir:
        # Create a test file
        test_file = os.path.join(tmpdir, "input.docx")
        from docx import Document
        doc = Document()
        doc.add_paragraph("original")
        doc.save(test_file)

        code = """
# Verify FILE_1 is available and points to the real file
import os
assert os.path.exists(FILE_1), f"FILE_1 does not exist: {FILE_1}"
from docx import Document
doc = Document(FILE_1)
assert doc.paragraphs[0].text == "original", f"Unexpected content: {doc.paragraphs[0].text}"
print("FILE_1 verification passed")
"""
        result = execute_code(code, output_dir=tmpdir, upload_dir=tmpdir,
                            file_paths=[test_file])
        assert result['success'], f"FILE_1 test failed: {result['stderr']}"


# ── Run ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import pytest
    pytest.main([__file__, "-v", "-s"])
