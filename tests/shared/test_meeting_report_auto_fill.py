#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.

import os
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from typing import Dict, List, Any
import requests
import logging

# 配置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class LLMMeetingProcessor:
    def __init__(self, llm_api_key: str = None, llm_base_url: str = None):
        """
        初始化LLM处理器
        支持OpenAI API兼容的接口
        """
        self.llm_api_key = llm_api_key or os.getenv('LLM_API_KEY')
        self.llm_base_url = llm_base_url or os.getenv('LLM_BASE_URL', 'https://api.deepseek.com/v1')

    def parse_meeting_with_llm(self, text_content: str) -> Dict[str, Any]:
        """
        使用LLM解析会议记录，提取结构化信息
        """
        prompt = f"""
            请从以下会议记录中提取结构化信息，并以JSON格式返回。提取的字段包括：
            
            必填字段：
            - meeting_theme: 会议主题
            - meeting_time: 会议时间（格式：YYYY年MM月DD日 HH:MM）
            - meeting_location: 会议地点
            - attendees: 参会人员列表
            - absentees: 缺席人员列表（如有）
            - presenter: 主持人
            - recorder: 记录人
            
            可选字段：
            - discussion_points: 讨论要点（列表形式）
            - decisions: 会议决议（列表形式）  
            - action_items: 行动项（列表形式，包含负责人和截止时间）
            - next_meeting: 下次会议安排
            - attachments: 附件列表
            
            会议记录内容：
            {text_content}
            
            请严格按照以下JSON格式返回，确保字段名称一致：
            {{
                "meeting_theme": "字符串",
                "meeting_time": "字符串", 
                "meeting_location": "字符串",
                "attendees": ["人员1", "人员2"],
                "absentees": ["人员1", "人员2"],
                "presenter": "字符串",
                "recorder": "字符串",
                "discussion_points": ["要点1", "要点2"],
                "decisions": ["决议1", "决议2"],
                "action_items": [
                    {{"item": "行动项描述", "assignee": "负责人", "deadline": "截止时间"}}
                ],
                "next_meeting": "字符串",
                "attachments": ["附件1", "附件2"]
            }}
        """

        try:
            response = self._call_llm_api(prompt)
            meeting_data = json.loads(response)
            return meeting_data
        except Exception as e:
            logger.error(f"LLM解析失败: {e}")
            return self._fallback_parsing(text_content)

    def _call_llm_api(self, prompt: str) -> str:
        """
        调用LLM API
        """
        headers = {
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {self.llm_api_key}'
        }

        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "system",
                 "content": "你是一个专业的会议记录分析助手，能够准确提取会议信息并输出结构化JSON数据。"},
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1
        }

        response = requests.post(
            f"{self.llm_base_url}/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )

        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        else:
            raise Exception(f"API调用失败: {response.status_code} - {response.text}")

    def _fallback_parsing(self, text_content: str) -> Dict[str, Any]:
        """
        LLM失败时的备用解析方案
        """
        # 基础正则匹配（简化版）
        data = {
            "meeting_theme": self._extract_by_pattern(text_content, r'主题[：:]\s*(.+)'),
            "meeting_time": self._extract_by_pattern(text_content, r'时间[：:]\s*(.+)'),
            "meeting_location": self._extract_by_pattern(text_content, r'地点[：:]\s*(.+)'),
            "attendees": self._extract_list(text_content, r'参会人员[：:]\s*(.+)'),
            "presenter": self._extract_by_pattern(text_content, r'主持人[：:]\s*(.+)'),
            "recorder": self._extract_by_pattern(text_content, r'记录人[：:]\s*(.+)'),
            "discussion_points": self._extract_multiline(text_content,
                                                         r'讨论[：:]\s*(.+?)(?=\n\s*\n|\n决议|\n行动项|$)'),
            "decisions": self._extract_multiline(text_content, r'决议[：:]\s*(.+?)(?=\n\s*\n|\n行动项|$)'),
        }

        return {k: v for k, v in data.items() if v}

    @staticmethod
    def _extract_by_pattern(text: str, pattern: str) -> str:
        match = re.search(pattern, text)
        return match.group(1).strip() if match else ""

    @staticmethod
    def _extract_list(text: str, pattern: str) -> List[str]:
        match = re.search(pattern, text)
        if match:
            return [item.strip() for item in match.group(1).split('、')]
        return []

    @staticmethod
    def _extract_multiline(text: str, pattern: str) -> List[str]:
        match = re.search(pattern, text, re.DOTALL)
        if match:
            lines = [line.strip() for line in match.group(1).split('\n') if line.strip()]
            return lines
        return []

class SmartTemplateFiller:
    def __init__(self):
        self.field_mapping = {
            # 中文字段映射
            '会议主题': 'meeting_theme',
            '会议名称': 'meeting_theme',
            '主题': 'meeting_theme',
            '会议时间': 'meeting_time',
            '时间': 'meeting_time',
            '会议地点': 'meeting_location',
            '地点': 'meeting_location',
            '参会人员': 'attendees',
            '出席人员': 'attendees',
            '缺席人员': 'absentees',
            '主持人': 'presenter',
            '记录人': 'recorder',
            '讨论内容': 'discussion_points',
            '讨论要点': 'discussion_points',
            '会议决议': 'decisions',
            '决议事项': 'decisions',
            '行动项': 'action_items',
            '下一步计划': 'action_items',
            '下次会议': 'next_meeting',
            '附件': 'attachments'
        }

    def analyze_template(self, template_path: str) -> List[Dict]:
        """
        分析模板结构，识别需要填充的字段
        """
        doc = Document(template_path)
        fields_found = []

        # 分析段落中的占位符
        for i, paragraph in enumerate(doc.paragraphs):
            fields = self._extract_fields_from_text(paragraph.text)
            if fields:
                fields_found.append({
                    'type': 'paragraph',
                    'index': i,
                    'text': paragraph.text,
                    'fields': fields,
                    'runs': [run.text for run in paragraph.runs]
                })

        # 分析表格中的占位符
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        fields = self._extract_fields_from_text(paragraph.text)
                        if fields:
                            fields_found.append({
                                'type': 'table_cell',
                                'table_index': table_idx,
                                'row_index': row_idx,
                                'cell_index': cell_idx,
                                'paragraph_index': para_idx,
                                'text': paragraph.text,
                                'fields': fields
                            })

        return fields_found

    def _extract_fields_from_text(self, text: str) -> List[str]:
        """
        从文本中提取可能的字段标识
        """
        fields = []

        # 匹配常见的字段标识模式
        patterns = [
            r'【(.+?)】',  # 【会议主题】
            r'《(.+?)》',  # 《会议时间》
            r'\{\{(.+?)\}\}',  # {{参会人员}}
            r'\[(.+?)\]',  # [主持人]
            r'(.+?)：',  # 会议主题：
            r'(.+?)\s*：',  # 会议主题 ：
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if self._is_likely_field(match):
                    fields.append(match.strip())

        return list(set(fields))  # 去重

    @staticmethod
    def _is_likely_field(text: str) -> bool:
        """
        判断文本是否可能是字段标识
        """
        field_keywords = [
            '会议', '主题', '时间', '地点', '人员', '主持', '记录',
            '讨论', '决议', '行动', '计划', '附件', '缺席', '出席'
        ]
        return any(keyword in text for keyword in field_keywords)

    def fill_template(self, template_path: str, meeting_data: Dict, output_path: str):
        """
        将会议数据填充到模板中
        """
        doc = Document(template_path)

        # 填充段落
        for paragraph in doc.paragraphs:
            self._fill_paragraph(paragraph, meeting_data)

        # 填充表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._fill_paragraph(paragraph, meeting_data)

        doc.save(output_path)
        logger.info(f"模板填充完成: {output_path}")

    def _fill_paragraph(self, paragraph, meeting_data: Dict):
        """
        填充单个段落
        """
        original_text = paragraph.text

        for field_display, field_key in self.field_mapping.items():
            if field_display in original_text and field_key in meeting_data:
                value = self._format_field_value(meeting_data[field_key])

                # 替换文本
                if field_display in paragraph.text:
                    # 保持原有格式
                    for run in paragraph.runs:
                        if field_display in run.text:
                            run.text = run.text.replace(field_display, value)

    @staticmethod
    def _format_field_value(value: Any) -> str:
        """
        格式化字段值
        """
        if isinstance(value, list):
            if all(isinstance(item, str) for item in value):
                return '、'.join(value)
            elif all(isinstance(item, dict) for item in value):
                # 处理行动项等复杂结构
                formatted = []
                for item in value:
                    if 'item' in item and 'assignee' in item:
                        formatted.append(
                            f"{item['item']}（负责人：{item['assignee']}，截止时间：{item.get('deadline', '待定')}）")
                    else:
                        formatted.append(str(item))
                return '；'.join(formatted)
            else:
                return '、'.join(str(item) for item in value)
        elif isinstance(value, dict):
            return str(value)
        else:
            return str(value) if value else "未填写"

class MeetingMinutesAutomation:
    def __init__(self, llm_api_key: str = None):
        self.llm_processor = LLMMeetingProcessor(llm_api_key)
        self.template_filler = SmartTemplateFiller()

    def process_meeting(self, txt_file_path: str, template_path: str, output_path: str = None):
        """
        主处理流程
        """
        # 1. 读取会议记录
        with open(txt_file_path, 'r', encoding='utf-8') as f:
            meeting_text = f.read()

        # 2. 使用LLM解析会议记录
        logger.info("正在使用LLM解析会议记录...")
        meeting_data = self.llm_processor.parse_meeting_with_llm(meeting_text)
        logger.info(f"解析结果: {json.dumps(meeting_data, ensure_ascii=False, indent=2)}")

        # 3. 分析模板结构
        logger.info("正在分析模板结构...")
        template_fields = self.template_filler.analyze_template(template_path)
        logger.info(f"发现字段: {template_fields}")

        # 4. 生成输出文件路径
        if not output_path:
            base_name = os.path.splitext(os.path.basename(txt_file_path))[0]
            output_path = f"filled_{base_name}.docx"

        # 5. 填充模板
        logger.info("正在填充模板...")
        self.template_filler.fill_template(template_path, meeting_data, output_path)

        return output_path, meeting_data

# 使用示例
def main():
    # 初始化自动化处理器
    processor = MeetingMinutesAutomation(llm_api_key="your_llm_api_key")

    # 处理单个会议记录
    output_file, meeting_data = processor.process_meeting(
        txt_file_path="meeting_notes.txt",
        template_path="meeting_template.docx",
        output_path="正式会议纪要.docx"
    )

    print(f"生成的会议纪要: {output_file}")
    print("会议数据:", json.dumps(meeting_data, ensure_ascii=False, indent=2))

# 批量处理
def batch_process():
    processor = MeetingMinutesAutomation(llm_api_key="your_llm_api_key")

    input_folder = "meeting_notes"
    template_folder = "templates"
    output_folder = "output"

    for filename in os.listdir(input_folder):
        if filename.endswith('.txt'):
            txt_path = os.path.join(input_folder, filename)
            template_path = os.path.join(template_folder, "standard_template.docx")
            output_path = os.path.join(output_folder, f"filled_{filename.replace('.txt', '.docx')}")

            try:
                processor.process_meeting(txt_path, template_path, output_path)
                print(f"成功处理: {filename}")
            except Exception as e:
                print(f"处理失败 {filename}: {e}")

if __name__ == "__main__":
    main()