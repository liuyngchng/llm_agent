#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.

"""
pip install python-docx cairosvg
# 安装常用的中文字体包
sudo apt update
sudo apt install fonts-wqy-microhei fonts-wqy-zenhei fonts-noto-cjk

# 安装字体工具
sudo apt install fontconfig

# 刷新字体缓存
fc-cache -fv

pip install python-docx cairosvg
"""
from docx import Document
from mermaid import Mermaid
import tempfile
import os


class MermaidDocxGenerator:
    def __init__(self):
        self.doc = Document()

    def add_mermaid_chart(self, title, mermaid_code, width_inches=5):
        """添加 Mermaid 图表到文档"""
        # 添加标题
        if title:
            self.doc.add_heading(title, level=2)

        try:
            # 生成图表
            mermaid = Mermaid(mermaid_code)

            # 创建临时文件
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                temp_path = temp_file.name

            # 保存为图片
            mermaid.save_as_png(temp_path)

            # 插入到文档
            from docx.shared import Inches
            self.doc.add_picture(temp_path, width=Inches(width_inches))

            # 清理
            os.unlink(temp_path)

            return True

        except Exception as e:
            self.doc.add_paragraph(f'错误: {str(e)}')
            self.doc.add_paragraph(f'代码: {mermaid_code}')
            return False

    def save(self, filename):
        """保存文档"""
        self.doc.save(filename)


if __name__ == '__main__':
    # 使用示例
    generator = MermaidDocxGenerator()

    # 添加图表
    generator.add_mermaid_chart(
        "系统架构图",
        """
    graph TB
        A[客户端] --> B[API网关]
        B --> C[用户服务]
        B --> D[订单服务]
        B --> E[支付服务]
        C --> F[(用户数据库)]
        D --> G[(订单数据库)]
        E --> H[(支付数据库)]
    """
    )

    generator.add_mermaid_chart(
        "开发流程",
        """
    graph LR
        A[需求] --> B[设计]
        B --> C[编码]
        C --> D[测试]
        D --> E[部署]
        E --> F[监控]
        F --> A
    """
    )

    generator.save('advanced_mermaid_example.docx')

