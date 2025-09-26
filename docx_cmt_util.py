#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.

"""
pip install python-docx
通过zipfile组件处理 docx 文档的批注信息，根据段落ID获取段落内容等
通过获取段落的批注，然后根据段落批注修改段落内容
"""
import logging.config
import os
import time
import zipfile
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import RGBColor, Cm

import docx_meta_util
from sys_init import init_yml_cfg
from agt_util import gen_txt
from docx_util import get_catalogue, refresh_current_heading, get_reference_from_vdb
from vdb_util import search_txt

logging.config.fileConfig('logging.conf', encoding="utf-8")
logger = logging.getLogger(__name__)

UPLOAD_FOLDER = 'upload_doc'

AI_GEN_TAG="[_AI生成_]"


def get_comments_dict(target_doc: str) -> dict:
    """
    获取文档中所有批注及其关联的段落ID
    返回格式: {
        comment_id: {
            "author": 作者,
            "date": 日期,
            "text": 批注内容,
            "paragraph_id": 段落ID  # 新增
        }
    }
    """
    if not os.path.exists(target_doc):
        logger.error(f"文件不存在: {target_doc}")
        return {}
    comments_dict = {}
    try:
        with zipfile.ZipFile(target_doc) as z:
            # 1. 解析批注 (comments.xml)
            if 'word/comments.xml' not in z.namelist():
                logger.warning("文档中没有批注")
                return {}
            with z.open('word/comments.xml') as f:
                comments_xml = ET.fromstring(f.read())
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                for comment in comments_xml.findall('.//w:comment', namespaces):
                    comment_id = comment.get(f'{{{namespaces["w"]}}}id')
                    if not comment_id:
                        continue

                    author = comment.get(f'{{{namespaces["w"]}}}author', '未知作者')
                    date = comment.get(f'{{{namespaces["w"]}}}date', '未知日期')
                    for t in comment.findall('.//w:t', namespaces):

                        if t.text:
                            text = ' '.join(t.text.strip())
                        else:
                            text = ""
                    comments_dict[comment_id] = {
                        "author": author,
                        "date": date,
                        "text": text,
                        "paragraph_id": None  # 初始化为None，后续填充
                    }

                    # 2. 解析正文 (document.xml)，关联批注和段落ID
                    if 'word/document.xml' not in z.namelist():
                        logger.warning("无法解析正文内容")

            with z.open('word/document.xml') as f:
                doc_xml = ET.fromstring(f.read())
                paragraphs = doc_xml.findall('.//w:p', namespaces)

                # 为每个段落生成唯一ID（如果没有现成的ID，用索引代替）
                for para_idx, paragraph in enumerate(paragraphs):
                    # 查找当前段落中的批注引用
                    comment_refs = paragraph.findall('.//w:commentReference', namespaces)
                    if not comment_refs:
                        continue

                    # 关联批注ID和段落ID（这里用索引作为段落ID）
                    for ref in comment_refs:
                        ref_id = ref.get(f'{{{namespaces["w"]}}}id')
                        if ref_id in comments_dict:
                            comments_dict[ref_id]["paragraph_id"] = para_idx
    except Exception as e:
        logger.error(f"解析文档时出错: {str(e)}", exc_info=True)
    return comments_dict


def get_paragraph_by_id(target_doc: str, paragraph_id: int) -> str:
    """
    通过段落ID获取段落文本
    :param target_doc: Word文档路径
    :param paragraph_id: 段落ID（索引或实际ID）
    :return: 段落文本（如未找到返回空字符串）
    """
    if not os.path.exists(target_doc):
        logger.error(f"文件不存在: {target_doc}")
        return ""

    try:
        with zipfile.ZipFile(target_doc) as z:
            if 'word/document.xml' not in z.namelist():
                logger.warning("无法解析正文内容")
                return ""

            with z.open('word/document.xml') as f:
                doc_xml = ET.fromstring(f.read())
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                paragraphs = doc_xml.findall('.//w:p', namespaces)

                if paragraph_id < 0 or paragraph_id >= len(paragraphs):
                    logger.warning(f"段落ID {paragraph_id} 超出范围")
                    return ""

                # 提取段落文本
                paragraph = paragraphs[paragraph_id]
                text = ' '.join(
                    t.text.strip()
                    for t in paragraph.findall('.//w:t', namespaces)
                    if t.text and t.text.strip()
                )
                return text

    except Exception as e:
        logger.error(f"解析段落时出错: {str(e)}", exc_info=True)
        return ""

def inspect_docx_structure():
    target_doc = "/home/rd/doc/文档生成/comment_test.docx"
    with zipfile.ZipFile(target_doc) as z:
        logger.info(f"文档包含的文件:{z.namelist()}")
        if 'word/comments.xml' not in z.namelist():
            logger.error("no word/comments.xml file")
            return
        with z.open('word/comments.xml') as f:
            logger.info(f"comments.xml内容:{f.read().decode('utf-8')}")

def get_para_comment_dict(file_fullpath: str) -> dict:
    """
    获取文档中所有修改批注及其关联的段落ID
    返回格式: {
        comment_id: {
            "author": 作者,
            "date": 日期,
            "text": 批注内容,
            "paragraph_id": 段落ID  # 新增
        }
    """

    comments_dict = get_comments_dict(file_fullpath)
    para_comments = {}
    for cid, data in comments_dict.items():
        para_id = data["paragraph_id"]
        if not para_id:
            continue
        para_comments[para_id] = data["text"]  # 一个段落只保留最后一条批注
    return para_comments

def test_get_comment():
    my_file = "/home/rd/doc/文档生成/comment_test.docx"
    comments_dict = get_comments_dict(my_file)
    logger.info(f"comments_dict={comments_dict}")
    for id, comment in comments_dict.items():
        logger.info(f"id:{id}, comment: {comment}")
        para_id = comment.get('paragraph_id')
        para_txt = get_paragraph_by_id(my_file, para_id)
        logger.info(f"para_id:{para_id}, para_txt:{para_txt}")

def modify_para_with_comment_prompt_in_process(task_id:int,
    target_doc: str, doc_ctx: str, comments_dict: dict, vdb_dir:str, cfg: dict, output_file_name:str):
    """
    将批注内容替换到对应段落
    :param task_id: 执行任务的ID
    :param target_doc: 需要修改的文档路径
    :param doc_ctx: 文档写作的背景信息
    :param comments_dict: 段落ID和段落批注的对应关系字典
    :param vdb_dir: 向量数据库的目录
    :param cfg: 系统配置，用于使用大模型的能力
    :param output_file_name: 输出文档的文件名
    """
    if not os.path.exists(target_doc):
        logger.error(f"输入文件不存在: {target_doc}")
        return
    if not comments_dict:
        logger.warning(f"文件批注信息为空")
        return
    logger.info(f"comments: {comments_dict}")
    doc = Document(target_doc)
    current_heading = []
    total_paragraphs = len(doc.paragraphs)
    gen_txt_count = 0
    comment_count = 0
    err_record = []
    for para_idx, para in enumerate(doc.paragraphs):
        percent = para_idx / total_paragraphs * 100
        process_percent_bar_info = (f"正在处理第 {para_idx + 1}/{total_paragraphs} 段文本，已识别 {comment_count} 个批注，"
            f"已生成 {gen_txt_count} 段文本，{get_elapsed_time(task_id)}")
        logger.info(f"{process_percent_bar_info}, 进度 {percent:.1f}%")
        try:
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, process_percent_bar_info, percent)
            refresh_current_heading(para, current_heading)
            if para_idx not in comments_dict:
                continue
            comment_count += 1
            logger.info(f"matched_comment_for_para_idx {para_idx}")
            comment_text = comments_dict[para_idx]
            catalogue = get_catalogue(target_doc)
            reference = get_reference_from_vdb(comment_text, vdb_dir, cfg['api'])
            modified_txt = gen_txt(doc_ctx, reference, comment_text, catalogue, str(current_heading), cfg)
            if modified_txt:
                gen_txt_count += 1
                para.clear()
                para.paragraph_format.first_line_indent = Cm(1) # set a first-line indent of approximately 1 cm (about 2 Chinese characters width)
                run = para.add_run(f"{AI_GEN_TAG}{modified_txt}")
                run.font.color.rgb = RGBColor(0, 0, 0)
                doc.save(output_file_name)
            else:
                logger.error(f"no_gen_txt_for_para, {para_idx}, comment {comment_text}")
        except Exception as e:
            err_info = f"在处理文档批注时发生异常: {str(e)}， 文档任务ID {task_id},错误信息 {str(e)}"
            err_record.append(err_info)
            logger.error(err_info, exc_info=True)
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, err_info)
            continue
    doc.save(output_file_name)
    docx_meta_util.save_docx_output_file_path_by_task_id(task_id, output_file_name)
    txt_info = f"任务已完成，共处理 {total_paragraphs} 段文本，识别 {comment_count} 个批注, 生成 {gen_txt_count} 段文本,发生错误 {len(err_record)} 处，错误信息 {err_record}"
    docx_meta_util.update_docx_file_process_info_by_task_id(task_id, txt_info, 100)

def modify_para_with_comment_prompt(target_doc: str,
        doc_ctx: str, comments_dict: dict, cfg: dict) -> Document:
    """
    将批注内容替换到对应段落，并将新文本设为红色
    :param target_doc: 需要修改的文档路径
    :param doc_ctx: 文档写作的大背景信息
    :param comments_dict: 段落ID和段落批注的对应关系字典
    :param cfg: 系统配置，用于使用大模型的能力
    """
    if not os.path.exists(target_doc):
        logger.error(f"输入文件不存在: {target_doc}")
        return
    if not comments_dict:
        logger.warning(f"文件批注信息为空")
        return
    logger.info(f"comments: {comments_dict}")
    doc = Document(target_doc)
    current_heading = []
    try:
        for para_idx, para in enumerate(doc.paragraphs):
            from docx_util import refresh_current_heading
            refresh_current_heading(para, current_heading)
            if para_idx not in comments_dict:
                continue
            logger.info(f"matched_comment_for_para_idx {para_idx}")
            comment_text = comments_dict[para_idx]
            catalogue = get_catalogue(target_doc)
            modified_txt = gen_txt(doc_ctx, "", comment_text, catalogue, str(current_heading), cfg)
            if modified_txt:
                para.clear()
                para.paragraph_format.first_line_indent = Cm(1) # set a first-line indent of approximately 1 cm (about 2 Chinese characters width)
                run = para.add_run(f"{AI_GEN_TAG}{modified_txt}")
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                logger.error(f"no_gen_txt_for_para, {para_idx}, comment {comment_text}")
    except Exception as e:
        logger.error(f"替换失败: {str(e)}", exc_info=True)
    return doc

def get_elapsed_time(start_timestamp: int) -> str:
    """
    计算任务处理时间
    :param start_timestamp: 任务开始时间戳
    """
    current_time = int(time.time())
    elapsed_seconds = current_time - start_timestamp
    minutes = elapsed_seconds // 60
    seconds = elapsed_seconds % 60
    return f"用时 {minutes}分{seconds}秒"

def test_modify_para_with_comment():
    my_cfg = init_yml_cfg()
    input_file = "/home/rd/doc/文档生成/comment_test.docx"
    para_comment_dict = get_para_comment_dict(input_file)
    doc_ctx = "我正在写一个可行性研究报告"
    output_doc = modify_para_with_comment_prompt(input_file, doc_ctx, para_comment_dict, my_cfg)
    output_file = "modify_comment_test.docx"
    output_doc.save(output_file)
    logger.info(f"处理完成，输出文件: {output_file}")

if __name__ == "__main__":
    # test_get_comment()
    # input_file = "/home/rd/doc/文档生成/comment_test.docx"
    # para_comment_dict = get_para_comment_dict(input_file)
    # logger.info(f"para_comment_dict: {para_comment_dict}")
    test_modify_para_with_comment()

