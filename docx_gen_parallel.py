#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.
import json
import logging.config
import multiprocessing
import os
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import List, Dict, Any
import threading

import psutil
from docx import Document
from docx.shared import RGBColor, Cm

import cfg_util
import docx_meta_util
import mermaid_render
from agt_util import gen_txt
from docx_util import get_elapsed_time, get_reference_from_vdb, is_3rd_heading, is_txt_para, \
    refresh_current_heading
from mermaid_render import MermaidRenderer

logging.config.fileConfig('logging.conf', encoding="utf-8")
logger = logging.getLogger(__name__)


class DocxGenerator:
    """
    DOCX文档生成器，使用多线程并行生成文本内容
    """

    def __init__(self, max_workers=None, timeout=300, consider_memory=True):
        """
        :param max_workers: 固定工作线程数，None则自动计算
        :param timeout: 任务超时时间
        :param consider_memory: 是否考虑内存使用情况
        """
        self.consider_memory = consider_memory

        if max_workers is None:
            max_workers = self._calculate_dynamic_workers()

        self.max_workers = max_workers
        self.executor = ThreadPoolExecutor(max_workers=max_workers)
        self.timeout = timeout
        self.lock = threading.Lock()
        self.start_time = None

        logger.info(f"初始化文档生成器 - 工作线程: {self.max_workers}, "
                    f"CPU核心: {multiprocessing.cpu_count()}, "
                    f"内存: {self._get_memory_info()}")

    def _calculate_dynamic_workers(self) -> int:
        """
        基于系统资源动态计算工作线程数
        考虑CPU核心数、内存使用情况等
        """
        try:
            cpu_count = multiprocessing.cpu_count()
            memory_info = DocxGenerator._get_memory_usage()
            base_workers = cpu_count * 2
            if self.consider_memory and memory_info['available_percent'] < 20:
                memory_factor = 0.5
                logger.warning(f"系统内存紧张({memory_info['available_percent']:.1f}%)，减少工作线程数")
            elif self.consider_memory and memory_info['available_percent'] < 40:
                memory_factor = 0.8
            else:
                memory_factor = 1.0
            optimal_workers = int(base_workers * memory_factor)
            optimal_workers = max(2, min(optimal_workers, 16))
            logger.info(f"动态计算工作线程 - CPU: {cpu_count}, "
                        f"内存可用: {memory_info['available_percent']:.1f}%, "
                        f"最终线程数: {optimal_workers}")
            return optimal_workers
        except Exception as e:
            logger.warning(f"动态计算工作线程失败，使用默认值: {e}")
            return 4

    @staticmethod
    def _get_memory_usage() -> dict:
        """
        获取系统内存使用情况
        """
        try:
            memory = psutil.virtual_memory()
            return {
                'total': memory.total,
                'available': memory.available,
                'used': memory.used,
                'available_percent': (memory.available / memory.total) * 100
            }
        except Exception as e:
            logger.warning(f"获取内存信息失败: {e}")
            return {'available_percent': 100}  # 默认认为内存充足

    @staticmethod
    def _get_memory_info() -> str:
        """获取内存信息字符串"""
        try:
            memory = DocxGenerator._get_memory_usage()
            return f"{memory['available_percent']:.1f}%可用"
        except:
            return "未知"

    def fill_doc_with_prompt_in_parallel(self, task_id: int, doc_ctx: str, target_doc: str,
                                         target_doc_catalogue: str, vdb_dir: str,
                                         sys_cfg: dict, output_file_name: str) -> str:
        """
        并行填充word文档
        """
        start_time = time.time() * 1000
        doc = Document(target_doc)
        try:
            logger.info(f"{task_id}, 开始处理文档 {target_doc}")
            tasks = DocxGenerator._collect_doc_with_prompt_gen_tasks(task_id, doc, doc_ctx, target_doc_catalogue,
                                                                     vdb_dir, sys_cfg)
            if not tasks:
                final_info = f"未检测到需要生成的文本段落"
                logger.info(f"{task_id}, {final_info}, {target_doc}")
                docx_meta_util.update_docx_file_process_info_by_task_id(task_id, final_info, 100)
                doc.save(output_file_name)
                return final_info

            initial_info = f"需处理 {len(tasks)} 个段落，启动 {self.executor._max_workers} 个任务"
            logger.info(f"{task_id}, {initial_info}")
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, initial_info, 0)
            doc_gen_results = self._exec_tasks(tasks, task_id, start_time, len(tasks), include_mermaid=True)
            self._insert_gen_para_to_doc(doc, doc_gen_results)

            # 保存文档
            doc.save(output_file_name)
            logger.info(f"{task_id}, 保存文档完成，{output_file_name}，耗时 {get_elapsed_time(start_time)}")
            docx_meta_util.save_docx_output_file_path_by_task_id(task_id, output_file_name)

            # 计算详细统计信息
            success_count = len([r for r in doc_gen_results.values() if r.get('success')])
            failed_count = len(tasks) - success_count


            # 处理Mermaid图表
            img_count = 0
            try:
                logger.info(f"{task_id}, 开始处理文档中的Mermaid图表")
                current_info = docx_meta_util.get_docx_info_by_task_id(task_id)
                docx_meta_util.update_docx_file_process_info_by_task_id(
                    task_id,
                    f"{current_info[0]['process_info']}, 开始处理文档配图",
                    95
                )
                mermaid_process_info = DocxGenerator._process_mermaid_in_document(
                    task_id,
                    output_file_name,
                    sys_cfg['api']['mermaid_api_uri'],
                    doc_gen_results
                )
                img_count = mermaid_process_info.get('img_count', 0)
                docx_meta_util.update_img_count_by_task_id(task_id, img_count)
                if not mermaid_process_info['success']:
                    failed_count += 1
                logger.info(f"{task_id}, Mermaid图表处理完成, {json.dumps(mermaid_process_info, ensure_ascii=False)}")
            except Exception as e:
                failed_count += 1
                logger.error(f"{task_id}, Mermaid图表处理失败: {str(e)}")
            total_time = get_elapsed_time(start_time)
            final_info = (f"文档处理完成，共执行 {len(tasks)} 个文本生成任务，"
                          f"成功生成 {success_count} 段文本和 {img_count} 张配图，失败任务 {failed_count} 个，{total_time}")

            if failed_count > 0:
                final_info += f"，失败任务的原因可在日志中查看具体的失败原因"
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, final_info, 100)
            logger.info(f"{task_id}, {final_info}，输出文件: {output_file_name}")
            return final_info

        except Exception as e:
            error_info = f"文档生成过程出现异常: {str(e)}"
            logger.error(f"{task_id}, {error_info}")
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, error_info, 100)
            try:
                doc.save(output_file_name)
            except Exception:
                pass
            return error_info

    @staticmethod
    def _collect_doc_with_prompt_gen_tasks(task_id: int, doc: Document, doc_ctx: str,
                                           target_doc_catalogue: str, vdb_dir: str,
                                           sys_cfg: dict) -> List[Dict[str, Any]]:
        """
        收集含有描述性的文本的docx文档的所有需要生成文本的任务信息
        """
        logger.info(f"{task_id}, start_collect_task")
        tasks = []
        current_heading = []
        para_count = len(doc.paragraphs)
        for index, para in enumerate(doc.paragraphs):
            refresh_current_heading(para, current_heading)
            check_if_is_txt_para = is_txt_para(para, current_heading, sys_cfg)
            if not check_if_is_txt_para:
                logger.info(f"{task_id}, 跳过非描述性的文本段落 {para.text}")
                continue
            task = {
                'unique_key': f"para_{len(tasks)}",
                'write_context': doc_ctx,
                'paragraph_prompt': para.text,
                'user_comment': "",
                'catalogue': target_doc_catalogue,
                'current_sub_title': current_heading[0] if current_heading else "",
                'current_heading': current_heading.copy(),
                'sys_cfg': sys_cfg,
                'vdb_dir': vdb_dir,
                'original_para': para,
            }
            tasks.append(task)
            process_info = f"正在处理第 {index}/{para_count}  段文本， 已创建 {len(tasks)} 个处理任务"
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, process_info)
        logger.info(f"_collect_task, {len(tasks)}")
        return tasks

    def _exec_tasks(self, tasks: List[Dict[str, Any]],
                    task_id: int, start_time: float,
                    total_tasks: int, include_mermaid: bool = False) -> Dict[str, Dict]:
        """
        并行执行文本生成任务
        :param tasks: 任务列表
        :param task_id: 任务ID
        :param start_time: 开始时间
        :param total_tasks: 总任务数
        :param include_mermaid: 是否包含Mermaid处理任务
        :return 执行完任务的结果
        """
        results = {}
        completed = 0
        future_to_key = {}
        last_update_time = time.time()
        update_interval = 2  # 每2秒更新一次进度，避免过于频繁

        # 如果包含Mermaid任务，总任务数需要加1
        actual_total_tasks = total_tasks + 1 if include_mermaid else total_tasks

        # 提交所有任务到线程池
        for task in tasks:
            future = self.executor.submit(DocxGenerator._gen_single_doc_paragraph, task)
            future_to_key[future] = task['unique_key']

        # 监控任务进度并收集结果
        for future in as_completed(future_to_key, timeout=self.timeout):
            key = future_to_key[future]
            try:
                result = future.result()
                results[key] = result
            except Exception as e:
                logger.error(f"{task_id}, 段落生成失败 {key}: {str(e)}, {task_id}")
                results[key] = {
                    'success': False,
                    'error': str(e),
                    'original_task': tasks[int(key.split('_')[1])]
                }

            completed += 1
            current_time = time.time()
            if current_time - last_update_time >= update_interval or completed == total_tasks:
                # 计算进度百分比时考虑Mermaid任务
                percent = int(completed / actual_total_tasks * 100)
                elapsed_time = get_elapsed_time(start_time)

                # 计算预估剩余时间
                if completed > 0:
                    elapsed_seconds = (time.time() * 1000 - start_time) / 1000
                    avg_time_per_task = elapsed_seconds / completed
                    remaining_tasks = actual_total_tasks - completed
                    estimated_remaining = avg_time_per_task * remaining_tasks

                    if estimated_remaining < 60:
                        remaining_str = f"约 {int(estimated_remaining)} 秒"
                    else:
                        remaining_str = f"约 {int(estimated_remaining / 60)} 分 {int(estimated_remaining % 60)} 秒"

                    progress_info = f"正在处理第 {completed}/{actual_total_tasks} 个任务，{elapsed_time}，剩余{remaining_str}"
                else:
                    progress_info = f"正在处理第 {completed}/{actual_total_tasks} 个任务，{elapsed_time}"

                docx_meta_util.update_docx_file_process_info_by_task_id(
                    task_id, progress_info, percent)
                last_update_time = current_time

        return results


    @staticmethod
    def _gen_single_doc_paragraph(task: Dict[str, Any]) -> Dict:
        """
        单个段落文本的生成任务（支持Mermaid图表）
        """
        try:
            # 获取参考文本
            references = get_reference_from_vdb(
                task['paragraph_prompt'],
                task['vdb_dir'],
                task['sys_cfg']['api']
            )
            logger.debug(f"gen_txt_user_comment, {task['user_comment']}")
            # 生成文本
            llm_txt = gen_txt(
                write_context=task['write_context'],
                references=references,
                paragraph_prompt=task['paragraph_prompt'],
                catalogue=task['catalogue'],
                current_sub_title=task['current_sub_title'],
                user_comment=task['user_comment'],
                cfg=task['sys_cfg']
            )
            word_count = len(llm_txt)


            return {
                'success': True,
                'generated_text': f"{cfg_util.AI_GEN_TAG}{llm_txt}",
                'original_para': task['original_para'],
                'current_heading': task['current_heading'],
                'contains_mermaid': '<mermaid>' in llm_txt,
                'word_count': word_count,
            }
        except Exception as e:
            heading_info = task['current_heading']
            logger.exception(f"生成段落失败: {str(e)}, single_task_args, {heading_info}")
            raise


    @staticmethod
    def _process_mermaid_in_document(task_id:int, doc_path: str, mermaid_api_uri: str, results: Dict[str, Dict]) -> Dict[str, Any]:
        """
        处理文档中的Mermaid图表
        :return: 处理结果信息
        """
        try:
            # 检查是否有包含Mermaid的内容
            has_mermaid = any(result.get('contains_mermaid') for result in results.values() if result.get('success'))
            mermaid_count = sum(1 for result in results.values() if result.get('success') and result.get('contains_mermaid'))
            if has_mermaid:
                logger.info(f"{task_id}, 检测到文档包含Mermaid图表 {mermaid_count}，开始处理: {doc_path}")
                mermaid_instance = MermaidRenderer(kroki_url=mermaid_api_uri)
                img_count = mermaid_instance.batch_process_mermaid_in_docx(task_id, doc_path)
                logger.info(f"{task_id}, 文档包含Mermaid图表处理完成: {doc_path}")
                return {
                    'success': True,
                    'mermaid_count': mermaid_count,
                    'has_mermaid': True,
                    'img_count': img_count,
                }
            else:
                logger.info(f"{task_id}, 文档未包含Mermaid图表，跳过处理: {doc_path}")
                return {
                    'success': True,
                    'mermaid_count': 0,
                    'has_mermaid': False,
                    'img_count':0,
                }
        except Exception as e:
            logger.error(f"{task_id}, 处理Mermaid图表时发生异常: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'img_count': 0,
            }

    @staticmethod
    def _insert_gen_para_to_doc(doc: Document, results: Dict[str, Dict]):
        """
        将生成的结果插入到文档中
        """
        # 按原始段落顺序处理结果
        sorted_keys = sorted(results.keys(), key=lambda x: int(x.split('_')[1]))

        for key in sorted_keys:
            result = results[key]
            if not result.get('success'):
                continue

            original_para = result['original_para']
            generated_text = result['generated_text']

            # 创建新段落并插入
            new_para = doc.add_paragraph()
            new_para.paragraph_format.first_line_indent = Cm(1)
            red_run = new_para.add_run(generated_text)
            red_run.font.color.rgb = RGBColor(0, 0, 0)

            # 插入到原始段落后面
            original_para._p.addnext(new_para._p)

    def modify_para_with_comment_prompt_in_parallel(self, task_id: int, target_doc: str, catalogue:str,
                                                doc_ctx: str, comments_dict: dict,
                                                vdb_dir: str, cfg: dict,
                                                output_file_name: str) -> str:
        """
        并行处理带有批注的文档
        :param task_id: 执行任务的ID
        :param target_doc: 需要修改的文档路径
        :param doc_ctx: 文档写作的背景信息
        :param comments_dict: 段落ID和段落批注的对应关系字典
        :param vdb_dir: 向量数据库的目录
        :param cfg: 系统配置，用于使用大模型的能力
        :param output_file_name: 输出文档的文件名
        """
        if not os.path.exists(target_doc):
            error_info = f"输入文件不存在: {target_doc}"
            logger.error(error_info)
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, error_info, 100)
            return error_info

        if not comments_dict:
            warning_info = "文件里未找到批注信息"
            logger.warning(warning_info)
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, warning_info, 100)
            return warning_info

        start_time = time.time() * 1000
        doc = Document(target_doc)

        try:
            info = f"处理带批注的文档 {target_doc}，共找到 {len(comments_dict)} 个批注"
            logger.info(info)
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, info)
            # 收集批注处理任务
            tasks = DocxGenerator._collect_doc_with_comment_gen_tasks(task_id, doc, catalogue, doc_ctx, comments_dict, vdb_dir, cfg)

            if not tasks:
                final_info = "未找到有效的批注处理任务"
                logger.info(final_info)
                docx_meta_util.update_docx_file_process_info_by_task_id(task_id, final_info, 100)
                doc.save(output_file_name)
                return final_info

            initial_info = f"需处理 {len(tasks)} 个批注段落，启动 {self.executor._max_workers} 个任务"
            logger.info(initial_info)
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, initial_info, 0)
            doc_gen_results = self._exec_tasks(tasks, task_id, start_time, len(tasks), include_mermaid=True)
            # DocxGenerator._update_doc_with_comments(doc_gen_results)
            DocxGenerator._update_doc_with_comments_using_revisions(doc, doc_gen_results)
            # 保存前启用修订显示
            doc.settings.track_revisions = True
            doc.settings.show_revisions = True
            doc.save(output_file_name)
            logger.info(f"保存批注处理文档完成: {output_file_name}")
            docx_meta_util.save_docx_output_file_path_by_task_id(task_id, output_file_name)
            # 统计结果
            success_count = len([r for r in doc_gen_results.values() if r.get('success')])
            failed_count = len(tasks) - success_count

            # 处理Mermaid图表
            img_count = 0
            try:
                logger.info(f"{task_id}, 开始处理文档中的Mermaid图表")
                current_info = docx_meta_util.get_docx_info_by_task_id(task_id)
                docx_meta_util.update_docx_file_process_info_by_task_id(
                    task_id,
                    f"{current_info[0]['process_info']}, 开始处理文档配图",
                    95
                )
                mermaid_process_info = DocxGenerator._process_mermaid_in_document(
                    task_id,
                    output_file_name,
                    cfg['api']['mermaid_api_uri'],
                    doc_gen_results
                )
                img_count = mermaid_process_info.get('img_count', 0)
                docx_meta_util.update_img_count_by_task_id(task_id, img_count)
                if not mermaid_process_info['success']:
                    failed_count += 1
                logger.info(f"{task_id}, Mermaid图表处理完成")
            except Exception as e:
                failed_count += 1  # 将Mermaid处理失败计入总失败数
                logger.error(f"{task_id}, Mermaid图表处理失败: {str(e)}")
            total_time = get_elapsed_time(start_time)
            final_info = (f"批注文档处理完成，共处理 {len(tasks)} 个批注段落，"
                          f"成功生成 {success_count} 段文本和 {img_count} 张配图，失败 {failed_count} 段，{total_time}")
            if failed_count > 0:
                final_info += "，失败段落可在日志中查看详情"
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, final_info, 100)
            logger.info(f"{task_id}, {final_info}，输出文件: {output_file_name}")
            return final_info
        except Exception as e:
            error_info = f"批注文档处理过程出现异常: {str(e)}"
            logger.error(f"{task_id}, {error_info}")
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, error_info, 100)
            try:
                doc.save(output_file_name)
            except Exception:
                pass
            return error_info

    @staticmethod
    def _collect_doc_with_comment_gen_tasks(task_id: int , doc: Document, catalogue:str, doc_ctx: str,
                               comments_dict: dict, vdb_dir: str,
                               cfg: dict) -> List[Dict[str, Any]]:
        """
        收集含有批注的文档的并行处理任务
        """
        tasks = []
        current_heading = []
        para_count = len(doc.paragraphs)
        for para_idx, para in enumerate(doc.paragraphs):
            # 更新当前标题
            refresh_current_heading(para, current_heading)
            check_if_txt_para = is_txt_para(para, current_heading, cfg)
            if not check_if_txt_para:
                logger.debug(f"{task_id}, 跳过非文本段落: {para.text}")
                continue
            # 检查当前段落是否有批注
            if para_idx not in comments_dict:
                logger.debug(f"{task_id}, 跳过无批注段落: {para.text}")
                continue
            comment_text = comments_dict[para_idx]
            if not comment_text or not comment_text.strip():
                logger.debug(f"{task_id}, 跳过无批注内容段落: {para.text}")
                continue

            task = {
                'unique_key': f"comment_{para_idx}",
                'write_context': doc_ctx,
                'paragraph_prompt': para.text,
                'user_comment': comment_text,
                'catalogue': catalogue,
                'current_sub_title': current_heading[0] if current_heading else "",
                'current_heading': current_heading.copy(),
                'sys_cfg': cfg,
                'vdb_dir': vdb_dir,
                'original_para': para,
                'para_index': para_idx,
            }
            logger.debug(f"{task_id}, user_comment: {task['user_comment']}")
            tasks.append(task)
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, f"正在处理第 {para_idx}/{para_count} 段文本，已创建 {len(tasks)} 个处理任务")
        return tasks

    @staticmethod
    def _update_doc_with_comments_using_revisions(doc: Document, results: Dict[str, Dict], author_name: str = "AI assistant powered by richard"):
        """
        使用修订模式更新文档中的批注段落，并添加作者信息
        """
        # 设置文档作者信息
        doc.core_properties.author = author_name
        doc.core_properties.last_modified_by = author_name

        # 按段落索引排序处理
        sorted_keys = sorted(results.keys(), key=lambda x: int(x.split('_')[1]))

        # 启用修订跟踪
        doc.settings.track_revisions = True

        for key in sorted_keys:
            result = results[key]
            if not result.get('success'):
                continue
            original_para = result['original_para']
            generated_text = result['generated_text'].replace(cfg_util.AI_GEN_TAG, '').strip()
            original_text = original_para.text.strip()
            # 清空原段落内容
            original_para.clear()
            # 添加原始文本（标记为删除）
            if original_text:
                deleted_run = original_para.add_run(original_text)
                deleted_run.font.strike = True  # 删除线表示删除
                # 设置颜色为红色表示删除
                deleted_run.font.color.rgb = RGBColor(255, 0, 0)

            # 添加新文本（标记为插入）
            if generated_text:
                inserted_run = original_para.add_run(f"{cfg_util.AI_GEN_TAG}{generated_text}")
                # 设置颜色为绿色表示插入
                inserted_run.font.color.rgb = RGBColor(0, 176, 80)
            original_para.paragraph_format.first_line_indent = Cm(1)


    @staticmethod
    def _update_doc_with_comments(results: Dict[str, Dict]):
        """
        用生成的结果更新文档中的批注段落
        """
        # 按段落索引排序处理
        sorted_keys = sorted(results.keys(), key=lambda x: int(x.split('_')[1]))

        for key in sorted_keys:
            result = results[key]
            if not result.get('success'):
                continue

            original_para = result['original_para']
            generated_text = result['generated_text']

            # 清空原段落内容并添加生成文本
            original_para.clear()
            original_para.paragraph_format.first_line_indent = Cm(1)
            run = original_para.add_run(generated_text)
            run.font.color.rgb = RGBColor(0, 0, 0)

    def fill_doc_without_prompt_in_parallel(self, task_id: int, doc_ctx: str, target_doc: str,
                                        target_doc_catalogue: str, vdb_dir: str,
                                        sys_cfg: dict, output_file_name: str) -> str:
        """
        并行填充word文档（只有三级目录，无需段落提示词）
        :param task_id: 执行任务的ID
        :param doc_ctx: 文档写作背景信息
        :param target_doc: 需要写的文档三级目录
        :param vdb_dir: 向量数据库的目录
        :param sys_cfg: 系统配置信息
        :param target_doc_catalogue: 需要写的文档的三级目录文本信息
        :param output_file_name: 输出文档的文件名
        """
        start_time = time.time() * 1000
        doc = Document(target_doc)

        try:
            logger.info(f"开始处理无提示词文档 {target_doc}")
            # 收集三级标题任务
            tasks = []
            for para in doc.paragraphs:
                if not is_3rd_heading(para):
                    continue

                task = {
                    'unique_key': f"heading_{len(tasks)}",
                    'write_context': doc_ctx,
                    'paragraph_prompt': "",
                    'user_comment': "",
                    'catalogue': target_doc_catalogue,
                    'current_sub_title': para.text,
                    'current_heading': [para.text],
                    'sys_cfg': sys_cfg,
                    'vdb_dir': vdb_dir,
                    'original_para': para,

                }
                tasks.append(task)

            if not tasks:
                final_info = f"未找到三级标题，输出文档： {output_file_name}"
                logger.info(final_info)
                docx_meta_util.update_docx_file_process_info_by_task_id(task_id, final_info, 100)
                doc.save(output_file_name)
                return final_info

            initial_info = f"需处理 {len(tasks)} 个三级标题，启动 {self.executor._max_workers} 个任务"
            logger.info(initial_info)
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, initial_info, 0)
            results = self._exec_tasks(tasks, task_id, start_time, len(tasks), include_mermaid=True)
            DocxGenerator._insert_gen_para_to_doc(doc, results)
            doc.save(output_file_name)
            logger.info(f"保存无提示词文档完成: {output_file_name}")
            docx_meta_util.save_docx_output_file_path_by_task_id(task_id, output_file_name)
            success_count = len([r for r in results.values() if r.get('success')])
            failed_count = len(tasks) - success_count
            # 处理Mermaid图表
            img_count = 0
            try:
                logger.info(f"{task_id}, 开始处理文档中的Mermaid图表")
                current_info = docx_meta_util.get_docx_info_by_task_id(task_id)
                docx_meta_util.update_docx_file_process_info_by_task_id(
                    task_id,
                    f"{current_info[0]['process_info']}, 开始处理文档配图",
                    95
                )
                mermaid_process_info = DocxGenerator._process_mermaid_in_document(
                    task_id,
                    output_file_name,
                    sys_cfg['api']['mermaid_api_uri'],
                    results
                )
                img_count = mermaid_process_info.get('img_count', 0)
                docx_meta_util.update_img_count_by_task_id(task_id, img_count)
                if not mermaid_process_info['success']:
                    failed_count += 1
                logger.info(f"{task_id}, Mermaid图表处理完成, {json.dumps(mermaid_process_info, ensure_ascii=False)}")
            except Exception as e:
                failed_count += 1  # 将Mermaid处理失败计入总失败数
                logger.error(f"{task_id}, Mermaid图表处理失败: {str(e)}")
            total_time = get_elapsed_time(start_time)
            final_info = (f"文档处理完成，共执行 {len(tasks)} 个文本生成任务，"
                          f"成功生成 {success_count} 段文本和 {img_count} 张配图，失败 {failed_count} 段，{total_time}")

            if failed_count > 0:
                final_info += "，失败标题可在日志中查看详情"

            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, final_info, 100)
            logger.info(f"{task_id}, {final_info}，输出文件: {output_file_name}")
            return final_info
        except Exception as e:
            error_info = f"文档生成过程出现异常: {str(e)}"
            logger.error(f"{task_id}, {error_info}")
            docx_meta_util.update_docx_file_process_info_by_task_id(task_id, error_info, 100)
            try:
                doc.save(output_file_name)
            except Exception:
                pass
            return error_info


    def shutdown(self):
        """主动关闭线程池"""
        self.executor.shutdown(wait=True)
        logger.info("文档生成器线程池已关闭")

    def __del__(self):
        """清理线程池"""
        try:
            self.executor.shutdown(wait=False)
        except:
            pass

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.shutdown()
