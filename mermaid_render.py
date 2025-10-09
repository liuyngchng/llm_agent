#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.


import logging.config
import os
import re
import requests
import tempfile
from docx import Document
from docx.shared import Inches
from PIL import Image

import cfg_util

logging.config.fileConfig('logging.conf', encoding="utf-8")
logger = logging.getLogger(__name__)


class MermaidRenderer:
    """
    Mermaid图表渲染器，将Mermaid脚本转换为图片并插入Word文档
    """

    def __init__(self, kroki_url="http://localhost:8000"):
        self.kroki_url = kroki_url
        self.supported_formats = ['png', 'svg']

    def render_script_to_img(self, mermaid_script: str, output_format: str = 'png') -> bytes:
        """
        将Mermaid脚本渲染为图片字节流
        :param mermaid_script: Mermaid脚本
        :param output_format: 输出格式 ('png' 或 'svg')
        :return: 图片字节流
        """
        if output_format not in self.supported_formats:
            raise ValueError(f"不支持的格式: {output_format}, 支持: {self.supported_formats}")

        try:
            # 清理Mermaid脚本
            cleaned_script = MermaidRenderer._clean_mermaid_script(mermaid_script)

            # 发送到Kroki服务
            response = requests.post(
                f"{self.kroki_url}/{output_format}",
                data=cleaned_script,
                headers={"Content-Type": "text/plain"},
                timeout=30
            )

            if response.status_code == 200:
                return response.content
            else:
                error_msg = (f"Kroki渲染失败[kroki_mermaid_script_render_fail]: "
                             f"{response.status_code},{response.text}, "
                             f"cleaned_script:[{cleaned_script}], "
                             f"original_script:[{mermaid_script}]")
                logger.error(error_msg)
                raise Exception(error_msg)

        except requests.exceptions.RequestException as e:
            error_msg = f"连接Kroki服务失败: {str(e)}"
            logger.error(error_msg)
            raise Exception(error_msg)
        except Exception as e:
            error_msg = f"Mermaid渲染异常: {str(e)}"
            logger.error(error_msg)
            raise Exception(error_msg)

    @staticmethod
    def _clean_mermaid_script(mermaid_script: str) -> str:
        """
        清理Mermaid脚本，确保语法正确
        """
        # 移除<mermaid>标签
        cleaned = re.sub(r'<mermaid>|</mermaid>', '', mermaid_script).strip()

        # 确保脚本以正确的图表类型开头
        if not cleaned.startswith(('graph', 'sequenceDiagram', 'erDiagram', 'gantt', 'classDiagram', 'stateDiagram', 'flowchart')):
            # 尝试自动添加缺失的图表类型声明
            if '->>' in cleaned or '->' in cleaned:
                cleaned = f"sequenceDiagram\n{cleaned}"
            elif '||--o{' in cleaned or '}|--||' in cleaned:
                cleaned = f"erDiagram\n{cleaned}"
            elif '-->' in cleaned or '[*]' in cleaned:
                cleaned = f"stateDiagram-v2\n{cleaned}"
            else:
                cleaned = f"graph TD\n{cleaned}"

        if cleaned.startswith('sequenceDiagram'):
            # 更完善的JSON内容识别和转义
            def escape_message_content(match):
                arrow_part = match.group(1)
                content = match.group(2).strip()
                # 如果内容以{开头，说明是JSON，需要转义并加引号
                if content.startswith('{'):
                    # 转义内部引号
                    escaped_content = content.replace('"', '\\"')
                    return f'{arrow_part} "{escaped_content}"'
                return match.group(0)

            # 匹配箭头操作符和后面的消息内容
            cleaned = re.sub(r'(\w+->>?\w+:)(.*?)(?=\n\s*\w+->>|\Z)',
                             escape_message_content,
                             cleaned,
                             flags=re.DOTALL)

        return cleaned



    def insert_mermaid_to_docx(self, doc_path: str, mermaid_script: str,
                               img_format: str = 'png', width: float = 5.0) -> str:
        """
        将Mermaid图表插入到Word文档中
        :param doc_path: Word文档路径
        :param mermaid_script: Mermaid脚本
        :param img_format: 图片格式
        :param width: 图片宽度（英寸）
        :return: 处理后的文档路径
        """
        try:
            # 渲染Mermaid为图片
            image_data = self.render_script_to_img(mermaid_script, img_format)

            # 加载文档
            doc = Document(doc_path)

            # 创建临时图片文件
            with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{img_format}') as temp_file:
                temp_file.write(image_data)
                temp_path = temp_file.name

            # 在文档末尾插入图片
            pic_para = doc.add_paragraph()
            pic_para.alignment = 1  # 居中对齐
            pic_para.add_run().add_picture(temp_path, width=Inches(width))

            # 保存文档
            doc.save(doc_path)

            # 清理临时文件
            os.unlink(temp_path)

            logger.info(f"Mermaid图表已成功插入文档: {doc_path}")
            return doc_path

        except Exception as e:
            logger.error(f"插入Mermaid图表失败: {str(e)}")
            raise

    def batch_process_mermaid_in_docx(self, task_id: int, doc_path: str, img_format: str = 'png') -> int:
        """
        批量处理文档中的所有Mermaid脚本
        :param doc_path: Word文档路径
        :param img_format: mermaid脚本转换为图片的格式
        :return: 处理的图片总数量
        """
        try:
            doc = Document(doc_path)
            mermaid_scripts = []

            # 查找所有包含Mermaid脚本的段落
            for para in doc.paragraphs:
                mermaid_match = re.search(r'<mermaid>(.*?)</mermaid>', para.text, re.DOTALL)
                if mermaid_match:
                    mermaid_script = mermaid_match.group(1)
                    mermaid_scripts.append({
                        'paragraph': para,
                        'script': mermaid_script,
                        'full_text': para.text
                    })

            if not mermaid_scripts:
                logger.info(f"{task_id}, 文档中未找到Mermaid脚本")
                return 0
            img_count = len(mermaid_scripts)
            logger.info(f"{task_id}, 找到需要处理的Mermaid脚本 {img_count} 个")

            # 处理每个Mermaid脚本
            for i, item in enumerate(mermaid_scripts):
                try:
                    # 渲染图表
                    image_data = self.render_script_to_img(item['script'], img_format)
                    # 创建临时图片文件
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{img_format}') as temp_file:
                        temp_file.write(image_data)
                        temp_path = temp_file.name

                    # 在Mermaid脚本段落后面插入图片
                    parent = item['paragraph']._p.getparent()
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = 1          # 图片左右居中
                    run = pic_para.add_run()
                    # 读取图片原始尺寸
                    with Image.open(temp_path) as img:
                        width, height = img.size
                        aspect_ratio = width / height

                    # 根据宽高比决定使用宽度优先还是高度优先
                    max_width = Inches(5.0)
                    max_height = Inches(4.0)
                    if aspect_ratio > 1:  # 横向图
                        run.add_picture(temp_path, width=max_width)
                    else:  # 竖向图
                        run.add_picture(temp_path, height=max_height)

                    # 将图片段落插入到原段落后面
                    parent.insert(parent.index(item['paragraph']._p) + 1, pic_para._p)

                    # 再插入文本段落
                    text_para = doc.add_paragraph()
                    text_run = text_para.add_run(f"{cfg_util.AI_GEN_TAG} 图 {i + 1}")
                    # 设置文本居中和加粗
                    text_para.alignment = 1  # 居中对齐
                    text_run.bold = True  # 加粗
                    # 将文本段落插入到图片段落后面
                    parent.insert(parent.index(pic_para._p) + 1, text_para._p)

                    # 清理临时文件
                    os.unlink(temp_path)
                    original_text = item['paragraph'].text
                    cleaned_text = re.sub(r'<mermaid>.*?</mermaid>', '', original_text, flags=re.DOTALL)
                    item['paragraph'].text = cleaned_text.strip()
                    logger.info(f"{task_id}, 已处理第 {i + 1} 个Mermaid图表")
                except Exception as e:
                    logger.error(f"{task_id}, 处理第 {i + 1} 个Mermaid脚本失败: {str(e)}, mermaid_scrip_fail, {item['paragraph'].text}")
                    continue
            # 保存文档
            doc.save(doc_path)
            logger.info(f"{task_id}, Mermaid图表批量处理完成: {doc_path}")
            return img_count
        except Exception as e:
            logger.error(f"{task_id}, 批量处理Mermaid图表失败: {str(e)}")
            raise

# 全局渲染器实例
instance = MermaidRenderer()


def test_mermaid_render():
    """测试Mermaid渲染功能"""
    renderer = MermaidRenderer()

    # 测试ER图
    er_diagram = """
        erDiagram
            DEVICE_INFO ||--o{ REAL_TIME_DATA : "1:N"
            DEVICE_INFO {
                varchar device_id PK
                varchar device_name
                varchar device_type
                datetime install_date
            }
            REAL_TIME_DATA {
                bigint record_id PK
                varchar device_id FK
                numeric sensor_value
                datetime collect_time
            }
        """

    try:
        image_data = renderer.render_script_to_img(er_diagram, 'png')
        logger.info(f"Mermaid渲染测试成功，图片大小: {len(image_data)} 字节")

        # 保存测试图片
        with open('test_mermaid.png', 'wb') as f:
            f.write(image_data)
        logger.info("测试图片已保存为 test_mermaid.png")

    except Exception as e:
        logger.error(f"Mermaid渲染测试失败: {str(e)}")


if __name__ == "__main__":
    test_mermaid_render()