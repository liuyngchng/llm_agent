#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.

"""
pip install python-docx cairosvg
# 安装常用的中文字体包
sudo apt update
sudo apt install fonts-wqy-microhei fonts-wqy-zenhei fonts-noto-cjk

# 安装字体工具
sudo apt install fontconfig

# 刷新字体缓存
fc-cache -fv
"""
# !/usr/bin/env python3
# -*- coding: utf-8 -*-
# Copyright (c) [2025] [liuyngchng@hotmail.com] - All rights reserved.

"""
pip install python-docx cairosvg
"""
from docx import Document
from docx.shared import Inches
import cairosvg
import io
import os


def add_svg_to_docx(doc, svg_content, width=Inches(6)):
    """
    将SVG内容插入到Word文档

    Args:
        doc: Word文档对象
        svg_content: SVG代码字符串
        width: 图片宽度
    """
    try:
        # 将SVG转换为PNG字节流
        png_data = cairosvg.svg2png(
            bytestring=svg_content.encode('utf-8'),
            # 设置DPI提高清晰度
            dpi=300
        )

        # 创建图片流
        image_stream = io.BytesIO(png_data)

        # 添加图片到文档
        doc.add_picture(image_stream, width=width)

        # 添加段落间距
        doc.add_paragraph()

    except Exception as e:
        print(f"SVG转换失败: {e}")
        doc.add_paragraph("【图表加载失败】")


def get_ubuntu_fonts():
    """
    返回Ubuntu系统可用的中文字体列表
    """
    # Ubuntu 常见的中文字体
    chinese_fonts = [
        "WenQuanYi Micro Hei",  # 文泉驿微米黑
        "WenQuanYi Zen Hei",  # 文泉驿正黑
        "Noto Sans CJK SC",  # Google思源黑体
        "Noto Serif CJK SC",  # Google思源宋体
        "AR PL UMing CN",  # 文鼎明体
        "AR PL UKai CN",  # 文鼎楷体
        "DejaVu Sans",  # 西文字体
        "Arial",
        "Helvetica",
        "sans-serif"
    ]
    return ", ".join(chinese_fonts)


# 使用示例
def create_document_with_svg():
    # 创建文档
    doc = Document()

    # 添加标题
    doc.add_heading('用户关系信息系统架构图', level=1)

    # 获取Ubuntu字体
    ubuntu_fonts = get_ubuntu_fonts()

    # SVG代码 - 使用Ubuntu系统字体
    svg_code = f'''<svg width="600" height="400" xmlns="http://www.w3.org/2000/svg" style="background-color: white;">
  <!-- 简化的架构图 -->
  <rect x="50" y="50" width="500" height="300" fill="none" stroke="#333" stroke-width="2"/>

  <!-- 层次框 -->
  <rect x="100" y="80" width="400" height="40" fill="#e3f2fd" stroke="#2196f3"/>
  <text x="300" y="105" text-anchor="middle" font-size="14" font-family="{ubuntu_fonts}" font-weight="bold">用户界面层</text>

  <rect x="100" y="140" width="400" height="40" fill="#e8f5e9" stroke="#4caf50"/>
  <text x="300" y="165" text-anchor="middle" font-size="14" font-family="{ubuntu_fonts}" font-weight="bold">业务逻辑层</text>

  <rect x="100" y="200" width="400" height="40" fill="#fff3e0" stroke="#ff9800"/>
  <text x="300" y="225" text-anchor="middle" font-size="14" font-family="{ubuntu_fonts}" font-weight="bold">数据访问层</text>

  <rect x="100" y="260" width="400" height="40" fill="#fce4ec" stroke="#e91e63"/>
  <text x="300" y="285" text-anchor="middle" font-size="14" font-family="{ubuntu_fonts}" font-weight="bold">数据库层</text>

  <!-- 箭头 -->
  <line x1="300" y1="120" x2="300" y2="140" stroke="#666" stroke-width="2" marker-end="url(#arrow)"/>
  <line x1="300" y1="180" x2="300" y2="200" stroke="#666" stroke-width="2" marker-end="url(#arrow)"/>
  <line x1="300" y1="240" x2="300" y2="260" stroke="#666" stroke-width="2" marker-end="url(#arrow)"/>

  <defs>
    <marker id="arrow" markerWidth="10" markerHeight="7" refX="9" refY="3.5" orient="auto">
      <polygon points="0 0, 10 3.5, 0 7" fill="#666"/>
    </marker>
  </defs>
</svg>'''

    # 插入SVG
    add_svg_to_docx(doc, svg_code)

    # 添加说明文字
    doc.add_paragraph("系统架构说明：本系统采用分层架构设计，各层职责如下：")
    doc.add_paragraph("1. 用户界面层：负责用户交互和界面展示")
    doc.add_paragraph("2. 业务逻辑层：处理核心业务规则和流程")
    doc.add_paragraph("3. 数据访问层：封装数据持久化操作")
    doc.add_paragraph("4. 数据库层：存储系统数据")

    # 保存文档
    doc.save('系统架构图.docx')
    print("文档已保存为 '系统架构图.docx'")


def check_ubuntu_fonts():
    """
    检查Ubuntu系统上的字体安装情况
    """
    print("检查Ubuntu字体安装情况...")

    # 检查常见中文字体包
    font_packages = [
        "fonts-wqy-microhei",  # 文泉驿微米黑
        "fonts-wqy-zenhei",  # 文泉驿正黑
        "fonts-noto-cjk",  # Google思源字体
        "fonts-arphic-ukai",  # 文鼎楷体
        "fonts-arphic-uming",  # 文鼎明体
    ]

    print("建议安装的字体包:")
    for pkg in font_packages:
        print(f"  sudo apt install {pkg}")

    print("\n当前可用的字体:")
    import subprocess
    try:
        result = subprocess.run(['fc-list', ':lang=zh'], capture_output=True, text=True)
        if result.returncode == 0:
            fonts = set()
            for line in result.stdout.split('\n'):
                if ':' in line:
                    font_name = line.split(':')[1].strip()
                    fonts.add(font_name)
            for font in sorted(fonts)[:10]:  # 只显示前10个
                print(f"  {font}")
    except:
        print("  无法获取字体列表，请安装 fontconfig: sudo apt install fontconfig")


if __name__ == "__main__":
    # 检查字体
    check_ubuntu_fonts()
    print("\n" + "=" * 50 + "\n")

    # 创建文档
    create_document_with_svg()