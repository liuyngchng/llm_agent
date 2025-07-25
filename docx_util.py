#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
pip install python-docx
处理docx的段落、标题、标题下的文本、文档目录、
修改docx的文档内容、设定新增的docx文档的文本格式等
"""
import logging.config
import os
import re
import time

from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH
from vdb_util import search_txt
from txt_util import get_txt_in_dir_by_keywords, strip_prefix_no

from sys_init import init_yml_cfg
from agt_util import classify_txt, gen_txt

logging.config.fileConfig('logging.conf', encoding="utf-8")
logger = logging.getLogger(__name__)

MIN_PROMPT_LEN = 20


AI_GEN_TAG="[_AI生成_]"


def get_reference_from_vdb(keywords: str, vdb_dir: str, sys_cfg: dict) -> str:
    """
    获取vdb中与关键词相关的文本
    :param keywords: 关键词
    :param vdb_dir: 向量数据库目录
    :param sys_cfg: 系统配置
    :return: 文本
    """
    if "" != vdb_dir and os.path.exists(vdb_dir):
        reference = search_txt(keywords, vdb_dir, 0.2, sys_cfg, 2).strip()
    else:
        logger.warning(f"vdb_dir_not_exist: {vdb_dir}, get no references")
        reference = ""
    # logging.info(f"vdb_get_txt:\n{reference}\nby_search_{keywords}")
    return reference


def extract_catalogue(target_doc: str) -> str:
    """
    生成 docx 文档的三级目录清单
    :param target_doc: 目标文档
    :return: 目录文本
    """
    doc = Document(target_doc)
    catalogue_lines = []
    # 初始化各级标题计数器 [一级, 二级, 三级]
    level_counters = [0, 0, 0]
    for para in doc.paragraphs:
        style_name = para.style.name.lower()  # 统一转换为小写
        # 检查是否为标题（兼容中英文样式名）
        if not style_name.startswith(('heading', '标题')):
            continue
        # 提取标题级别数字
        level_str = ''.join(filter(str.isdigit, style_name))
        if not level_str:
            continue
        level = int(level_str)
        if level <1 or level > 3:
            continue
        # 仅处理1-3级标题
        # 更新计数器：当前级别+1，更低级别清零
        level_index = level - 1
        level_counters[level_index] += 1
        for i in range(level_index + 1, 3):
            level_counters[i] = 0

        # 生成编号 (如 "2.1.1")
        number_parts = []
        for i in range(level):
            number_parts.append(str(level_counters[i]))
        number_str = '.'.join(number_parts)

        # 添加缩进和目录行
        indent = "  " * (level - 1)
        catalogue_lines.append(f"{indent}{number_str} {para.text}")
    return "\n".join(catalogue_lines)

def refresh_current_heading(para: Paragraph, heading: list) -> bool:
    """
    更新当前标题，并返回是否为三级标题
    :param para: 段落
    :param heading: 当前标题
    :return: 是否为三级标题
    """
    if "Heading" not in para.style.name:
        return False

    level = int(para.style.name.split()[-1])  # 提取数字
    headings = {1: [], 2: [], 3: [], 4: [], 5: []}  # 按需扩展层级
    headings[level].append(para.text)
    logger.info(f"heading_part_caught: H{level}: {para.text}")
    if None == heading:
        logger.error("heading_list_cant't_be_refresh")
        return True
    if len(heading) != 0:
        heading.pop()
    heading.append(para.text)
    return True

def is_3rd_heading(para: Paragraph) -> bool:
    """
    判断是否为三级标题
    :param para: 段落
    :return: 是否为三级标题
    """
    if "Heading" not in para.style.name:
        return False

    level = int(para.style.name.split()[-1])  # 提取数字
    if level == 3:
        return True
    else:
        return False

def get_outline(file_name: str) -> str:
    """
    获取word文档的三级目录，输出格式如下：
    # 1. 一级目录
    ## 1.1 二级目录
    ### 1.1.1 三级目录
    """
    doc = Document(file_name)
    result_lines = []  # 存储每行目录文本
    chapter_num = 0    # 一级标题编号
    section_num = 0    # 二级标题编号
    subsection_num = 0 # 三级标题编号

    for para in doc.paragraphs:
        style_name = para.style.name.lower()

        # 跳过非标题段落
        if not style_name.startswith(('heading', '标题')):
            continue

        # 提取标题级别
        level_str = ''.join(filter(str.isdigit, style_name))
        if not level_str:
            continue
        level = int(level_str)

        # 只处理1-3级标题
        if level < 1 or level > 3:
            continue

        text = para.text.strip()
        if not text:
            continue

        # 处理一级标题
        if level == 1:
            chapter_num += 1
            section_num = 0
            subsection_num = 0
            result_lines.append(f"# {chapter_num}. {text}")

        # 处理二级标题
        elif level == 2 and chapter_num > 0:
            section_num += 1
            subsection_num = 0
            result_lines.append(f"## {chapter_num}.{section_num} {text}")

        # 处理三级标题
        elif level == 3 and chapter_num > 0 and section_num > 0:
            subsection_num += 1
            result_lines.append(f"### {chapter_num}.{section_num}.{subsection_num} {text}")

    return "\n".join(result_lines)

def is_prompt_para(para: Paragraph, current_heading:list, sys_cfg: dict) -> bool:
    """
    判断写作要求文档中的每段文本，是否为用户所提的写作要求文本
    :param para: 写作要求word文档中的一个段落
    :param current_heading: 当前para 所在的目录
    :param sys_cfg: 系统配置，涉及大模型的地址等
    return False： 不是写作要求； True： 是写作要求
    """

    pattern = r'^(图|表)\s*\d+[\.\-\s]'  # 匹配"图1."/"表2-"等开头
    labels = ["需要生成文本", "不需要生成文本"]
    if refresh_current_heading(para, current_heading):
        return False
    if "TOC" in para.style.name or para._element.xpath(".//w:instrText[contains(.,'TOC')]"):
        logger.info(f"doc_table_of_content: {para.text}")
        return False
    if "Caption" in para.style.name or re.match(pattern, para.text):
        logger.info(f"table_or_picture_title: {para.text}")
        return False
    if not para.text:
        return False
    if len(para.text.strip()) < MIN_PROMPT_LEN:
        logger.info(f"ignored_short_txt {para.text}")
        return False
    if len(current_heading) == 0 or len(current_heading[0]) < 2:
        logger.info(f"heading_err_for_para, {current_heading}, {para.text}")
        return False
    classify_result = classify_txt(labels, para.text, sys_cfg, True)
    if labels[1] in classify_result:
        # logger.debug(f"classify={classify_result}, tile={current_heading}, para={para.text}")
        return False
    # logger.debug(f"classify={classify_result}, tile={current_heading}, para={para.text}")
    return True

def fill_doc_without_prompt_in_progress(uid: int, task_id:str, thread_lock, progress_info:dict, doc_ctx: str, target_doc: str,
    target_doc_catalogue: str, vdb_dir: str, sys_cfg: dict, output_file_name:str):
    """
    :param uid: 用户ID
    :param task_id: 执行任务的ID，时间戳的整数字符串
    :param thread_lock: A thread lock
    :param progress_info: task process information dict with task_id as key
    :param doc_ctx: 文档写作背景信息
    :param target_doc: 需要写的文档三级目录，以及各个章节的具体写作需求
    :param vdb_dir: 向量数据库的目录
    :param sys_cfg: 系统配置信息
    :param target_doc_catalogue: 需要写的文档的三级目录文本信息
    :param output_file_name: 输出文档的文件名
    """
    doc = Document(target_doc)
    gen_txt_count = 0
    total_paragraphs = len(doc.paragraphs)
    txt_info = ""
    for index, my_para in enumerate(doc.paragraphs):
        percent = index / total_paragraphs * 100
        process_percent_bar_info = (f"正在处理第 {index+1}/{total_paragraphs} 段文本，"
            f"已生成 {gen_txt_count} 段文本，{get_elapsed_time(task_id)}，进度 {percent:.1f}%")
        logger.info(process_percent_bar_info)
        update_process_info(thread_lock, uid, task_id, progress_info, process_percent_bar_info, percent)
        try:
            hd_check = is_3rd_heading(my_para)
            if not hd_check:
                continue
            # logger.info(f"prompt_txt_of_heading {current_heading}, {my_para.text}")
            reference = get_reference_from_vdb(my_para.text, vdb_dir, sys_cfg['api'])
            llm_txt = gen_txt(doc_ctx, reference, "", target_doc_catalogue, my_para.text, sys_cfg, )
            gen_txt_count += 1
        except Exception as ex:
            logger.error("fill_doc_job_err_to_break", ex)
            update_process_info(thread_lock, uid, task_id, progress_info, f"在处理文档的过程中出现了异常，任务已中途退出")
            break
        new_para = doc.add_paragraph()
        new_para.paragraph_format.first_line_indent = Cm(1) # set a first-line indent of approximately 1 cm (about 2 Chinese characters width)
        red_run = new_para.add_run(llm_txt)
        red_run.font.color.rgb = RGBColor(0, 0, 0)
        my_para._p.addnext(new_para._p)
        doc.save(output_file_name)
        logger.info(f"fill_doc_job_success_save_doc, {output_file_name}")
        if gen_txt_count > 0:
            txt_info = f"任务已完成，共处理 {total_paragraphs} 段文本，已生成 {gen_txt_count} 段文本，进度 100%"
        else:
            txt_info = f"任务已完成，共处理 {total_paragraphs} 段文本，进度 100%，未检测到写作要求文本"
        update_process_info(thread_lock, uid, task_id, progress_info, txt_info, 100)
    logger.info(f"{txt_info}, 所有内容已输出至文件 {output_file_name}")

def fill_doc_with_prompt_in_progress(uid: int, task_id:str, thread_lock, process_info:dict, doc_ctx: str, target_doc: str,
    target_doc_catalogue: str, vdb_dir: str, sys_cfg: dict, output_file_name:str):
    """
    :param uid: 用户ID
    :param task_id: 执行任务的ID
    :param thread_lock: A thread lock
    :param process_info: task process information dict with task_id as key
    :param doc_ctx: 文档写作背景信息
    :param target_doc: 需要写的文档三级目录，以及各个章节的具体写作需求
    :param vdb_dir: 向量数据库的目录
    :param sys_cfg: 系统配置信息
    :param target_doc_catalogue: 需要写的文档的三级目录文本信息
    :param output_file_name: 输出文档的文件名
    """
    doc = Document(target_doc)
    gen_txt_count = 0
    current_heading = []
    total_paragraphs = len(doc.paragraphs)
    for index, my_para in enumerate(doc.paragraphs):
        percent = index / total_paragraphs * 100
        process_percent_bar_info = (f"正在处理第 {index+1}/{total_paragraphs} 段文本，"
            f"已生成 {gen_txt_count} 段文本，{get_elapsed_time(task_id)}, 进度 {percent:.1f}%")
        logger.info(process_percent_bar_info)
        update_process_info(thread_lock, uid, task_id, process_info, process_percent_bar_info, percent)
        try:
            is_prompt = is_prompt_para(my_para, current_heading, sys_cfg)
            if not is_prompt:
                continue
            # logger.info(f"prompt_txt_of_heading {current_heading}, {my_para.text}")
            reference = get_reference_from_vdb(my_para.text, vdb_dir, sys_cfg['api'])
            # with thread_lock:
            #     thread_lock[task_id] = f"正在处理文本[{my_para.text}]"
            llm_txt = gen_txt(doc_ctx, reference, my_para.text, target_doc_catalogue, current_heading[0], sys_cfg, )
            gen_txt_count += 1
            # with thread_lock:
            #     thread_lock[task_id] = f"生成文本：{llm_txt}"
        except Exception as ex:
            logger.error("fill_doc_job_err_to_break", ex)
            update_process_info(thread_lock, uid, task_id, process_info, f"在处理文档的过程中出现了异常，任务已中途退出", percent)
            break
        new_para = doc.add_paragraph()
        new_para.paragraph_format.first_line_indent = Cm(1) # set a first-line indent of approximately 1 cm (about 2 Chinese characters width)
        red_run = new_para.add_run(f"{AI_GEN_TAG}{llm_txt}")
        red_run.font.color.rgb = RGBColor(0, 0, 0)
        my_para._p.addnext(new_para._p)
        doc.save(output_file_name)
    txt_info = f"任务已完成，共处理 {total_paragraphs} 段文本，已生成 {gen_txt_count} 段文本，{get_elapsed_time(task_id)}，进度 100%"
    if gen_txt_count == 0:
        txt_info += f"未检测到创作需求描述，您可以尝试在需要创作的段落处填写： 描述/列出/简述xxxxx, 写作需求描述文字数量大于20个汉字"
    update_process_info(thread_lock, uid, task_id, process_info, txt_info, 100.0)
    logger.info(f"{txt_info}, 所有内容已输出至文件 {output_file_name}")


def fill_doc_with_prompt(doc_ctx: str, source_dir: str, target_doc: str, target_doc_catalogue: str, vdb_dir: str, sys_cfg: dict) -> Document:
    """
    :param doc_ctx: 文档写作背景信息
    :param source_dir: 提供的样本文档
    :param target_doc: 需要写的文档三级目录，以及各个章节的具体写作需求
    :param vdb_dir: 向量数据库的目录
    :param sys_cfg: 系统配置信息
    :param target_doc_catalogue: 需要写的文档的三级目录文本信息
    """
    doc = Document(target_doc)
    current_heading = []
    total_paragraphs = len(doc.paragraphs)
    for index, my_para in enumerate(doc.paragraphs):
        percent = (index + 1) / total_paragraphs * 100
        process_percent_bar_info = f"正在处理第 {index+1}/{total_paragraphs} 段文本，进度 {percent:.1f}%"
        logger.info(f"percent: {process_percent_bar_info}")
        try:
            is_prompt = is_prompt_para(my_para, current_heading, sys_cfg)
            if not is_prompt:
                continue
            # logger.info(f"prompt_txt_of_heading {current_heading}, {my_para.text}")
            reference = get_reference_from_vdb(my_para.text, vdb_dir, sys_cfg['api'])
            if source_dir and os.path.exists(source_dir):
                source_para_txt = get_txt_in_dir_by_keywords(strip_prefix_no(current_heading[0]), source_dir)
                demo_txt = f"{source_para_txt}\n{reference}"
                demo_txt = demo_txt.replace("\n", " ").strip()
            else:
                demo_txt = f"{reference}"

            llm_txt = gen_txt(doc_ctx, demo_txt, my_para.text, target_doc_catalogue, current_heading[0], sys_cfg, )
            logger.info(f"llm_txt_for_instruction[{my_para.text}]\n===gen_llm_txt===\n{llm_txt}")
        except Exception as ex:
            logger.error("fill_doc_job_err_to_break", ex)
            break
        # if len(my_txt) > 0:
        new_para = doc.add_paragraph()
        new_para.paragraph_format.first_line_indent = Cm(1) # set a first-line indent of approximately 1 cm (about 2 Chinese characters width)
        red_run = new_para.add_run(f"{AI_GEN_TAG}{llm_txt}")
        red_run.font.color.rgb = RGBColor(255, 0, 0)
        my_para._p.addnext(new_para._p)
    return doc

def update_process_info(thread_lock, uid: int, task_id: str, process_info: dict, txt_info: str, percent: float):
    """
    :param thread_lock: A thread lock
    :param uid： 用户ID
    :param task_id: 执行任务的ID
    :param process_info: task process information dict with uid and task_id as key,
        process_info = {
           "uid1": {
                "taskId1": {
                    "percent": 0.0,
                    "text": "目前的状况是个啥？",
                    "timestamp": time.time(),
                    "elapsed_time": "xxx分xxx秒",
                }
           },
           "uid2": {
                "taskId2": {
                    "percent": 0.0,
                    "text": "目前的状况是个啥？",
                    "timestamp": time.time(),
                    "elapsed_time": "xxx分xxx秒",
                }
           }
        }
    :param txt_info: 任务进度信息
    :param percent: 任务进度百分比
    """
    if not process_info:
        process_info = {}
    if process_info.get(uid, None) is None:
                process_info[uid] = {}
    if process_info[uid].get(task_id, None) is None:
        process_info[uid][task_id] = {}
    with thread_lock:
        if not txt_info:
            process_info[uid][task_id]["text"] = txt_info
        if not percent:
            process_info[uid][task_id]["percent"] = percent
        process_info[uid][task_id]["timestamp"] = time.time()
        process_info[uid][task_id]["elapsed_time"] = get_elapsed_time(task_id)

def get_process_info(uid: int, task_id: str, process_info: dict) -> dict:
    """
    获取文档当前处理进度信息
    :param task_id: 执行任务的ID
    :param process_info: task process information dict with task_id as key
        process_info = {
            "uid1": {
                "taskId1": {
                    "percent": 0.0,
                    "text": "目前的状况是个啥？",
                    "timestamp": time.time(),
                    "elapsed_time": "xxx分xxx秒",
                }
            },
            "uid2": {
                "taskId2": {
                    "percent": 0.0,
                    "text": "目前的状况是个啥？",
                    "timestamp": time.time(),
                    "elapsed_time": "xxx分xxx秒",
                }
           }
        }
    """
    process_info.get(task_id, {"text": "未知状态", "percent": 0, "timestamp": time.time(), "elapsed_time":""})
    info = {
        "task_id": task_id,
        "progress": process_info.get("text", ""),
        "percent": process_info.get("percent", ""),
        "elapsed_time": process_info.get("elapsed_time", "")
    }
    return info

def get_catalogue(target_doc: str) -> str:
    catalogue_file = "my_catalogue.txt"
    if os.path.exists(catalogue_file):
        logger.info(f"文件 {catalogue_file} 已存在")
        with open(catalogue_file, 'r', encoding='utf-8') as f:
            my_catalogue = f.read()
            logger.info(f"目录内容已从文件 {catalogue_file} 读取")
    else:
        logger.info(f"文件 {catalogue_file} 不存在，将创建")
        my_catalogue = extract_catalogue(target_doc)
        with open(catalogue_file, 'w', encoding='utf-8') as f:
            f.write(my_catalogue)
            logger.info(f"目录内容已写入 {catalogue_file}")
    return my_catalogue


def gen_docx_template_with_outline(task_id: str, os_dir:str, title: str, outline: str) -> str:
    """
    :param task_id: 执行任务的ID
    :param os_dir: 输出的本地文件目录
    :param title: 文档标题
    :param outline: 三级目录文本信息，结构如下\n1. 一级标题\n  1.1 二级标题\n    1.1.1 三级标题\n
    即outline中的各行文本中，一级标题前没有空格，二级标题前有2个空格，三级标题前有4个空格
    :return: 包含三级目录的Word docx 文档文件名称，只是文件名，不是fullpath
    """
    doc = Document()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    title_run.font.size = Pt(28)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加分页符使目录从新页面开始
    doc.add_page_break()
    lines = outline.strip().split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        indent_level = 0
        for char in line:
            if char == ' ':
                indent_level += 1
            elif not char.isspace():
                break
        # 添加标题并获取段落对象
        if indent_level == 0:
            p = doc.add_heading(stripped, level=1)
        elif indent_level == 2:
            p = doc.add_heading(stripped, level=2)
        else:
            p = doc.add_heading(stripped, level=3)
        # 统一设置格式：黑体、黑色
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')  # 中文字体
    # 保存文件
    filename = f"{task_id}_outline.docx"
    save_path = os.path.join(os_dir, filename)
    doc.save(save_path)
    return filename


def get_elapsed_time(start_timestamp: str) -> str:
    """
    计算任务处理时间
    :param start_timestamp: 任务开始时间戳
    """
    start_time = int(start_timestamp)
    current_time = int(time.time())
    elapsed_seconds = current_time - start_time
    minutes = elapsed_seconds // 60
    seconds = elapsed_seconds % 60
    return f"用时 {minutes}分{seconds}秒"

if __name__ == "__main__":
    my_cfg = init_yml_cfg()
    my_source_dir = "/home/rd/doc/文档生成/knowledge_base"
    # my_target_doc = "/home/rd/doc/文档生成/output_template.docx"
    my_target_doc = "/home/rd/doc/文档生成/2.docx"
    # test_catalogue = extract_catalogue(my_target_doc)
    # logger.info(f"doc_catalogue: {test_catalogue}")
    my_doc_ctx = "我正在写一个可行性研究报告"
    doc_catalogue = get_catalogue(my_target_doc)
    logger.info(f"my_target_doc_catalogue: {doc_catalogue}")
    my_vdb_dir = "./vdb/vdb_idx_332987902_26"
    output_doc = fill_doc_with_prompt(my_doc_ctx, my_source_dir, my_target_doc, doc_catalogue, my_vdb_dir, my_cfg)
    output_file = 'doc_output.docx'
    output_doc.save(output_file)
    logger.info(f"save_content_to_file: {output_file}")
